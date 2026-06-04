"""生成 Playwright 测试用的招标文档"""
import os
import sys

# 尝试安装 python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx --quiet")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_test_docx(output_path: str, filename: str = "test_bidding_doc.docx"):
    """生成一份测试用招标文档（含常见合规问题）"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)

    doc.add_heading('第一章 招标公告', level=1)
    doc.add_paragraph('本采购项目采用公开招标方式，欢迎合格供应商投标。预算金额500万元。')
    doc.add_paragraph('项目编号：BHG-2026-001')
    doc.add_paragraph('采购人：某市政府采购中心')

    doc.add_heading('第二章 招标范围', level=1)
    doc.add_paragraph('本次采购内容包括XXX系统建设及运维服务。')
    doc.add_paragraph('交货地点：采购人指定地点。')
    doc.add_paragraph('交货期限：合同签订后30日内。')

    doc.add_heading('第三章 投标人资格要求', level=1)
    doc.add_paragraph('1. 投标人应具有独立承担民事责任的能力。')
    doc.add_paragraph('2. 投标人必须为本市注册企业，注册资本不低于1000万元。')
    doc.add_paragraph('3. 本项目不接受联合体投标。')
    doc.add_paragraph('4. 投标人应具有ISO9001质量管理体系认证。')
    doc.add_paragraph('5. 投标人近三年内无重大违法记录。')

    doc.add_heading('第四章 评审办法', level=1)
    doc.add_paragraph('本项目采用综合评分法。')
    doc.add_paragraph('技术方案：40分')
    doc.add_paragraph('价格分：30分')
    doc.add_paragraph('业绩分：30分')
    doc.add_paragraph('总分100分，得分最高者中标。')

    doc.add_heading('第五章 投标须知', level=1)
    doc.add_paragraph('投标截止时间：2026年7月1日9:00。')
    doc.add_paragraph('投标有效期：90天。')
    doc.add_paragraph('投标保证金：人民币10万元整。')
    doc.add_paragraph('开标时间：2026年7月1日9:00。')
    doc.add_paragraph('开标地点：某市公共资源交易中心。')

    full_path = os.path.join(output_path, filename)
    doc.save(full_path)
    print(f"✅ 测试文档已创建: {full_path}")
    return full_path


if __name__ == '__main__':
    output = sys.argv[1] if len(sys.argv) > 1 else '/tmp'
    create_test_docx(output)
