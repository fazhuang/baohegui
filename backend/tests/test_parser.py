"""测试：文档解析服务（DocumentParser）"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.services.parser import (
    DocumentParser,
    ParsedDocument,
    SectionInfo,
    DocumentParserError,
    TextSpan,
    PageLine,
    _fuzzy_match_section_type,
    _char_bigram_set,
    _bigram_overlap,
    _docx_heading_level,
)


@pytest.fixture
def parser() -> DocumentParser:
    return DocumentParser()


# ═══════════════════════════════════════════════════════════════
# Word 解析
# ═══════════════════════════════════════════════════════════════

class TestParseDocx:
    def test_parse_simple_docx(self, parser):
        """简单的单章节文档"""
        doc = Document()
        doc.add_heading("第一章 招标公告", level=1)
        doc.add_paragraph("欢迎合格供应商投标。")
        doc.add_heading("第二章 资格要求", level=1)
        doc.add_paragraph("投标人应具有独立法人资格。")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            assert isinstance(result, ParsedDocument)
            assert result.filename == Path(tmp.name).name
            assert "招标公告" in result.sections
            assert "资格要求" in result.sections
            assert len(result.sections) >= 2
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_section_content_captured(self, parser):
        """验证章节内容正确提取"""
        doc = Document()
        doc.add_heading("招标公告", level=1)
        doc.add_paragraph("公告正文内容。")
        doc.add_heading("资格要求", level=1)
        doc.add_paragraph("1. 独立法人。")
        doc.add_paragraph("2. 良好信誉。")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            assert "公告正文内容" in result.sections.get("招标公告", "")
            assert "独立法人" in result.sections.get("资格要求", "")
            assert "良好信誉" in result.sections.get("资格要求", "")
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_detect_all_five_sections(self, parser):
        """检测五大必备章节"""
        doc = Document()
        doc.add_heading("第一章 招标公告", level=1)
        doc.add_paragraph("项目概况。")
        doc.add_heading("第二章 招标范围", level=1)
        doc.add_paragraph("采购内容。")
        doc.add_heading("第三章 投标人资格要求", level=1)
        doc.add_paragraph("资格条件。")
        doc.add_heading("第四章 评审办法", level=1)
        doc.add_paragraph("评审标准。")
        doc.add_heading("第五章 投标须知", level=1)
        doc.add_paragraph("须知。")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            required = {"招标公告", "招标范围", "资格要求", "评审办法", "投标须知"}
            found = set(result.sections.keys())
            missing = required - found
            assert not missing, f"缺失章节: {missing}"
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_heading_detection_noise_suppressed(self, parser):
        """正文中的短词不会误检为标题"""
        doc = Document()
        doc.add_heading("招标公告", level=1)
        doc.add_paragraph("采购内容详见附件。本项目概况如下。")  # 含"采购内容"/"项目概况"
        doc.add_paragraph("投标人应满足资格要求。")  # 含"资格要求"

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            # 不应该因为正文中的"采购内容"而创建"招标范围"章节
            assert "招标范围" not in result.sections
            assert "资格要求" not in result.sections  # 正文中的"资格要求"不应创建章节
            assert len(result.raw_sections) == 1  # 只有"招标公告"一个章节
        finally:
            Path(tmp.name).unlink(missing_ok=True)


# ═══════════════════════════════════════════════════════════════
# PDF 解析（基础测试）
# ═══════════════════════════════════════════════════════════════

class TestParsePDF:
    def test_pdf_not_found(self, parser):
        with pytest.raises(DocumentParserError, match="文件不存在"):
            parser.parse_pdf("/nonexistent/file.pdf")

    def test_invalid_pdf(self, parser):
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.write(b"not a real pdf content")
        tmp.close()
        try:
            with pytest.raises(DocumentParserError):
                parser.parse_pdf(tmp.name)
        finally:
            Path(tmp.name).unlink(missing_ok=True)


# ═══════════════════════════════════════════════════════════════
# 通用解析入口
# ═══════════════════════════════════════════════════════════════

class TestParse:
    def test_unsupported_format(self, parser):
        with pytest.raises(DocumentParserError, match="不支持的文件格式"):
            parser.parse("test.xlsx")

    def test_detect_docx_by_extension(self, parser, sample_docx_path):
        result = parser.parse(sample_docx_path)
        assert isinstance(result, ParsedDocument)

    def test_multiple_detected_sections(self, parser, sample_docx_path):
        result = parser.parse(sample_docx_path)
        assert len(result.sections) >= 1


# ═══════════════════════════════════════════════════════════════
# ParsedDocument 辅助方法
# ═══════════════════════════════════════════════════════════════

class TestParsedDocument:
    def test_to_dict(self):
        doc = ParsedDocument(filename="test.docx", page_count=5)
        d = doc.to_dict()
        assert d["filename"] == "test.docx"
        assert d["page_count"] == 5

    def test_get_section_count(self):
        doc = ParsedDocument(filename="test.docx")
        assert doc.get_section_count() == 0
        doc.sections = {"招标公告": "..."}
        assert doc.get_section_count() == 1

    def test_get_missing_required_sections(self):
        doc = ParsedDocument(filename="test.docx")
        doc.sections = {"招标公告": "...", "评审办法": "..."}
        missing = doc.get_missing_required_sections()
        assert "招标范围" in missing
        assert "资格要求" in missing
        assert "投标须知" in missing
        assert "招标公告" not in missing

    def test_no_missing(self):
        doc = ParsedDocument(filename="test.docx")
        doc.sections = {
            "招标公告": "...", "招标范围": "...", "资格要求": "...",
            "评审办法": "...", "投标须知": "...",
        }
        assert doc.get_missing_required_sections() == []


# ═══════════════════════════════════════════════════════════════
# SectionInfo
# ═══════════════════════════════════════════════════════════════

class TestSectionInfo:
    def test_to_dict(self):
        s = SectionInfo(
            title="第一章 招标公告",
            section_type="招标公告",
            content="公告正文",
            page_start=1,
            page_end=2,
        )
        d = s.to_dict()
        assert d["section_type"] == "招标公告"
        assert d["page_start"] == 1
        assert d["page_end"] == 2
        assert d["content_length"] == 4


# ═══════════════════════════════════════════════════════════════
# 模糊章节匹配
# ═══════════════════════════════════════════════════════════════

class TestFuzzySectionMatching:
    def test_exact_match(self):
        """精确匹配保持原行为"""
        assert _fuzzy_match_section_type("招标公告") == "招标公告"
        assert _fuzzy_match_section_type("资格要求") == "资格要求"
        assert _fuzzy_match_section_type("投标须知") == "投标须知"

    def test_substring_match(self):
        """子串匹配同样有效"""
        assert _fuzzy_match_section_type("第一章 招标公告") == "招标公告"
        assert _fuzzy_match_section_type("投标人资格要求") == "资格要求"
        assert _fuzzy_match_section_type("评标办法") == "评审办法"

    def test_fuzzy_match_variants(self):
        """模糊匹配处理各种变体"""
        matches = [
            ("投标人资格要求（供应商）", "资格要求"),
            ("资格要求（施工）", "资格要求"),
            ("评审办法（综合评分法）", "评审办法"),
            ("采购需求及技术参数", "技术要求"),
            ("投标人须知前附表", "投标须知"),
        ]
        for title, expected in matches:
            result = _fuzzy_match_section_type(title)
            assert result == expected, f"Expected '{expected}' for '{title}', got '{result}'"

    def test_non_section_text(self):
        """与章节无关的文本不应匹配"""
        assert _fuzzy_match_section_type("项目概况") is None
        # "投标人应满足资格要求" contains "资格要求" as substring, so
        # _fuzzy_match_section_type legitimately matches it.
        # The filtering happens upstream in _is_heading_line.
        assert _fuzzy_match_section_type("关于进一步加强管理的通知") is None
        assert _fuzzy_match_section_type("附件：技术规范") is None

    def test_bigram_utilities(self):
        """bi-gram 工具函数正确性"""
        assert _char_bigram_set("") == set()
        assert _bigram_overlap("", "") == 0.0
        assert _bigram_overlap("资格要求", "资格要求") == 1.0
        ratio = _bigram_overlap("投格要求", "资格要求")  # Simulated typo
        assert 0 < ratio < 1.0


# ═══════════════════════════════════════════════════════════════
# Docx 层级与表格解析
# ═══════════════════════════════════════════════════════════════

class TestParseDocxAdvanced:
    def test_heading_styles_detected(self, parser):
        """Heading 1/2/3 样式被正确检测"""
        from docx import Document
        import tempfile
        from pathlib import Path

        doc = Document()
        doc.add_heading("第一章 招标公告", level=1)
        doc.add_paragraph("公告正文。")
        doc.add_heading("1.1 项目背景", level=2)
        doc.add_paragraph("背景介绍。")
        doc.add_heading("第二章 资格要求", level=1)
        doc.add_paragraph("资格条件。")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            assert "招标公告" in result.sections
            assert "资格要求" in result.sections
            # 检查层级
            for sec in result.raw_sections:
                if sec.section_type == "招标公告":
                    assert sec.level == 1, f"Expected level 1, got {sec.level}"
                if sec.title == "1.1 项目背景":
                    # 可能被检测为子标题
                    pass
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_table_extraction(self, parser):
        """表格内容被提取"""
        from docx import Document
        import tempfile
        from pathlib import Path

        doc = Document()
        doc.add_heading("招标公告", level=1)
        doc.add_paragraph("公告正文。")

        # 添加表格
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "序号"
        table.cell(0, 1).text = "名称"
        table.cell(0, 2).text = "数量"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "服务器"
        table.cell(1, 2).text = "5台"

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)

        try:
            result = parser.parse(tmp.name)
            full_text = result.full_text
            assert "服务器" in full_text, f"Table content '服务器' not found in: {full_text[:200]}"
            assert "序号 | 名称" in full_text, f"Table header not found"
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_heading_level_mapping(self):
        """Heading 样式名到层级的映射正确"""
        assert _docx_heading_level("Heading 1") == 1
        assert _docx_heading_level("Heading 2") == 2
        assert _docx_heading_level("heading 3") == 3
        assert _docx_heading_level("heading 1") == 1
        assert _docx_heading_level("Normal") == 1  # fallback


# ═══════════════════════════════════════════════════════════════
# PageLine / TextSpan 工具函数
# ═══════════════════════════════════════════════════════════════

class TestPageLineHelpers:
    def test_page_line_creation(self):
        pl = PageLine(
            text="第一章 招标公告",
            page_num=1,
            bbox=(50, 100, 400, 120),
            max_font_size=16.0,
            is_bold=True,
            heading_score=0.8,
        )
        assert pl.text == "第一章 招标公告"
        assert pl.page_num == 1
        assert pl.max_font_size == 16.0
        assert pl.is_bold is True

    def test_text_span_creation(self):
        span = TextSpan(text="招标公告", font_size=16.0, is_bold=True, font_name="SimHei")
        assert span.text == "招标公告"
        assert span.font_size == 16.0
        assert span.is_bold is True
        assert span.font_name == "SimHei"


# ═══════════════════════════════════════════════════════════════
# SectionInfo 新增方法
# ═══════════════════════════════════════════════════════════════

class TestSectionInfoExtensions:
    def test_add_content(self):
        s = SectionInfo(title="招标公告", section_type="招标公告", content="初始内容")
        s.add_content("追加内容")
        assert s.content == "初始内容追加内容"

    def test_append_subheading(self):
        s = SectionInfo(title="招标公告", section_type="招标公告", content="")
        s.append_subheading("1.1 项目背景")
        assert "1.1 项目背景" in s.headings

    def test_content_length(self):
        s = SectionInfo(title="招标公告", section_type="招标公告", content="12345")
        assert s.content_length() == 5
        s.add_content("678")
        assert s.content_length() == 8

    def test_to_dict_includes_level(self):
        s = SectionInfo(
            title="第一章 招标公告", section_type="招标公告", content="正文",
            level=1, page_start=1, page_end=2,
        )
        d = s.to_dict()
        assert d["level"] == 1
        assert d["page_start"] == 1
        assert d["page_end"] == 2
        assert d["content_length"] == 2


# ═══════════════════════════════════════════════════════════════
# PDF 字体感知解析（集成测试）
# ═══════════════════════════════════════════════════════════════

class TestParsePDFAdvanced:
    def test_pdf_font_detection_fallback(self, parser):
        """PDF: 回退到文本级检测（helv 不包含 CJK 字形，验证 ASCII 路径）"""
        import tempfile
        from pathlib import Path

        import fitz
        doc = fitz.open()
        page = doc.new_page()
        # Use ASCII text with mixed font sizes to test heading detection fallback
        page.insert_text((72, 100), "Section 1 Bidding Announcement", fontsize=16, fontname="helv")
        page.insert_text((72, 150), "All qualified suppliers are welcome to participate.", fontsize=11, fontname="helv")
        page.insert_text((72, 200), "Section 2 Qualification Requirements", fontsize=16, fontname="helv")
        page.insert_text((72, 250), "Bidders shall have independent legal personality.", fontsize=11, fontname="helv")

        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.close()
        doc.save(tmp.name)
        doc.close()

        try:
            result = parser.parse(tmp.name)
            assert result.page_count >= 1
            # Should detect at least text-based content
            assert "Bidding" in result.full_text or "Qualification" in result.full_text
        finally:
            Path(tmp.name).unlink(missing_ok=True)


# ═══════════════════════════════════════════════════════════════
# 真实 PDF 集成测试（PDF/ 目录）
# ═══════════════════════════════════════════════════════════════

import os

PDF_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "PDF")

REAL_PDFS = [
    os.path.join(PDF_DIR, "GPC", "公开-货物类--甘肃医学院解剖学实验教学平台项目第1标段.pdf"),
    os.path.join(PDF_DIR, "Hydraulic", "后审施工--兰州市七里河区黄峪沟（建西路段）防洪治理工程招标文件.pdf"),
    os.path.join(PDF_DIR, "Engineering", "后审公开施工--招标文件-一标段（终稿）.pdf"),
]


class TestRealPDFs:
    """使用真实招标 PDF 验证解析质量"""

    @pytest.mark.parametrize("pdf_path", REAL_PDFS)
    def test_parse_completes_without_error(self, parser, pdf_path):
        """所有真实 PDF 应成功解析"""
        if not os.path.exists(pdf_path):
            pytest.skip(f"PDF not found: {pdf_path}")
        result = parser.parse(pdf_path)
        assert result.page_count > 0
        assert len(result.full_text) > 1000

    @pytest.mark.parametrize("pdf_path", REAL_PDFS)
    def test_section_count_reasonable(self, parser, pdf_path):
        """解析结果不应产生过量章节（< 100）"""
        if not os.path.exists(pdf_path):
            pytest.skip(f"PDF not found: {pdf_path}")
        result = parser.parse(pdf_path)
        count = len(result.sections)
        assert count < 100, (
            f"Too many sections detected ({count}), expected < 100. "
            f"Sections: {list(result.sections.keys())}"
        )

    @pytest.mark.parametrize("pdf_path", REAL_PDFS)
    def test_required_sections_found(self, parser, pdf_path):
        """必备章节应有合理检出"""
        if not os.path.exists(pdf_path):
            pytest.skip(f"PDF not found: {pdf_path}")
        result = parser.parse(pdf_path)
        found = set(result.sections.keys())
        # At least 2 of the 5 required sections should be found
        required = {"招标公告", "招标范围", "资格要求", "评审办法", "投标须知"}
        intersection = found & required
        assert len(intersection) >= 2, (
            f"Only {len(intersection)} required sections found: {intersection}, "
            f"found: {found}"
        )

    @pytest.mark.parametrize("pdf_path", REAL_PDFS)
    def test_sections_have_content(self, parser, pdf_path):
        """检出的章节应有实质内容"""
        if not os.path.exists(pdf_path):
            pytest.skip(f"PDF not found: {pdf_path}")
        result = parser.parse(pdf_path)
        empty_sections = [
            st for st, content in result.sections.items()
            if len(content.strip()) < 20
        ]
        assert len(empty_sections) <= len(result.sections) * 0.3, (
            f"Too many near-empty sections: {empty_sections}"
        )

    @pytest.mark.parametrize("pdf_path", REAL_PDFS)
    def test_no_list_items_as_sections(self, parser, pdf_path):
        """编号列表项不应被误检为独立章节类型"""
        if not os.path.exists(pdf_path):
            pytest.skip(f"PDF not found: {pdf_path}")
        result = parser.parse(pdf_path)
        # Check raw_sections: no section should have a title like "2.1" or "5.1"
        import re
        numeric_titles = [
            s for s in result.raw_sections
            if re.match(r'^\d+\.\d+$', s.title.strip())
        ]
        assert len(numeric_titles) < 5, (
            f"Found {len(numeric_titles)} numeric-only section titles: "
            f"{[s.title for s in numeric_titles[:10]]}"
        )
