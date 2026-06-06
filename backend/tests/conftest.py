"""pytest 共享 Fixtures — 使用 SQLite 文件数据库，无需 PostgreSQL/MinIO/LLM API

核心策略：
- 文件级 SQLite 数据库（所有代码路径共享同一数据库）
- pytest_configure hook 在最早期阶段 mock MinIO 单例
- 自动创建/清理测试数据库
- LLM mock mode 始终开启
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock

import pytest

# ── 确定测试数据库路径（必须在导入任何 app 模块前创建）─────
_PROJ_TMP = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".test_tmp")
os.makedirs(_PROJ_TMP, exist_ok=True)
_TEST_DB_PATH = os.path.join(_PROJ_TMP, "test.db")

# ── 强制 mock / test 环境变量 ──────────────────────────────
os.environ["BHG_LLM_MOCK_MODE"] = "true"
os.environ["BHG_DEBUG"] = "true"
os.environ["BHG_SECRET_KEY"] = "test-secret-key---overriding-default-value-for-ci"
os.environ["BHG_DATABASE_URL"] = f"sqlite:///{_TEST_DB_PATH}"
os.environ["BHG_LOG_LEVEL"] = "error"
os.environ["BHG_MINIO_ENDPOINT"] = "localhost:9000"
os.environ["BHG_MINIO_ACCESS_KEY"] = "mock-access-key"
os.environ["BHG_MINIO_SECRET_KEY"] = "mock-secret-key"
os.environ["BHG_CORS_ORIGINS"] = "http://localhost:5173"


def _fake_download(object_key, target_path):
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_paragraph("公开招标公告正文。")
    doc.add_heading("第二章 招标范围", level=1)
    doc.add_paragraph("采购内容。")
    doc.add_heading("第三章 投标人资格要求", level=1)
    doc.add_paragraph("独立法人。指定品牌产品。")
    doc.add_heading("第四章 评审办法", level=1)
    doc.add_paragraph("综合评分法。")
    doc.add_heading("第五章 投标须知", level=1)
    doc.add_paragraph("投标截止时间2026年7月1日。投标保证金。")
    doc.save(target_path)
    return target_path


def _fake_local_path(storage_path):
    class _Ctx:
        def __enter__(self):
            if storage_path.startswith("uploads/"):
                tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                _fake_download(storage_path, tmp.name)
                tmp.close()
                self._p = tmp.name
                return tmp.name
            return storage_path

        def __exit__(self, *a):
            if hasattr(self, "_p"):
                Path(self._p).unlink(missing_ok=True)

    return _Ctx()


def pytest_configure(config):
    """
    在测试收集前 mock MinIO 单例。
    这是最早的可拦截点，确保 minio_service 在导入时已被 patch。
    """
    from app.services import minio_service as _ms

    _ms.minio_service.upload = MagicMock(return_value="uploads/test.docx")
    _ms.minio_service.download = MagicMock(side_effect=_fake_download)
    _ms.minio_service.delete = MagicMock(return_value=None)
    _ms.minio_service.ensure_bucket = MagicMock(return_value=None)
    _ms.minio_service.local_path = _fake_local_path
    _ms.minio_service._client = MagicMock()


def pytest_sessionfinish(session, exitstatus):
    """清理测试数据库目录"""
    import shutil

    if os.path.exists(_PROJ_TMP):
        shutil.rmtree(_PROJ_TMP, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════
# 数据库 fixtures
# ═══════════════════════════════════════════════════════════════

_TABLES_TO_CLEAN = [
    "compliance_reports",
    "document_sections",
    "uploaded_files",
    "user_quotas",
    "audit_logs",
    "feedback_records",
    "rule_confidences",
    "rule_mappings",
    "rule_versions",
    "rules",
    "announcements",
    "users",
]


def _ensure_tables(engine):
    """确保所有表已创建（幂等）"""
    from app.core.audit import AuditBase
    from app.models.announcement import Base as AnnouncementBase
    from app.models.document import Base as DocumentBase
    from app.models.rule import Base as RuleBase
    from app.models.subscription import Base as SubscriptionBase
    from app.services.feedback_service import FeedbackRecord, RuleConfidence  # noqa: F401

    for base in [DocumentBase, RuleBase, AuditBase, AnnouncementBase, SubscriptionBase]:
        base.metadata.create_all(bind=engine, checkfirst=True)


def _clean_all_tables(engine):
    """清空所有表数据（保留表结构）"""
    from sqlalchemy import text

    _ensure_tables(engine)
    with engine.begin() as conn:
        conn.execute(text("PRAGMA foreign_keys=OFF"))
        for table in _TABLES_TO_CLEAN:
            conn.execute(text(f"DELETE FROM {table}"))
        conn.execute(text("PRAGMA foreign_keys=ON"))


@pytest.fixture(scope="session")
def _test_db_engine():
    """Session 级别的 SQLite 引擎"""
    from sqlalchemy import create_engine

    engine = create_engine(
        f"sqlite:///{_TEST_DB_PATH}",
        connect_args={"check_same_thread": False},
        pool_pre_ping=True,
    )
    _ensure_tables(engine)
    return engine


@pytest.fixture(autouse=True)
def db_session(_test_db_engine):
    """每个测试获得全新干净的数据库会话"""
    _clean_all_tables(_test_db_engine)

    from sqlalchemy.orm import sessionmaker

    session_local = sessionmaker(bind=_test_db_engine)
    session = session_local()
    try:
        yield session
    finally:
        session.rollback()
        session.close()


# ═══════════════════════════════════════════════════════════════
# FastAPI TestClient fixtures
# ═══════════════════════════════════════════════════════════════


@pytest.fixture
def client(db_session):
    """FastAPI TestClient，数据库依赖注入为测试 session"""
    from fastapi.testclient import TestClient

    from app.db.database import get_db
    from app.main import app

    def _override_get_db():
        yield db_session

    app.dependency_overrides[get_db] = _override_get_db
    try:
        with TestClient(app) as tc:
            yield tc
    finally:
        app.dependency_overrides.pop(get_db, None)


@pytest.fixture
def auth_headers(db_session):
    """返回 admin 用户的认证头（自动种子 admin 用户，仅用于测试）"""
    from app.core.security import create_access_token, hash_password
    from app.models.user import User

    admin = User(
        username="admin",
        hashed_password=hash_password("admin123"),
        role="admin",
        company="测试",
        email="admin@test.com",
    )
    db_session.add(admin)
    db_session.commit()
    db_session.refresh(admin)

    token = create_access_token(user_id=admin.id, role="admin")
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def user_auth_headers(db_session):
    """返回普通用户的认证头（自动种子 user 用户，仅用于测试）"""
    from app.core.security import create_access_token, hash_password
    from app.models.user import User

    user = User(
        username="testuser",
        hashed_password=hash_password("user123"),
        role="user",
        company="测试单位",
        email="user@test.com",
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)

    token = create_access_token(user_id=user.id, role="user")
    return {"Authorization": f"Bearer {token}"}


# ═══════════════════════════════════════════════════════════════
# 原始 fixtures（保留给引擎级测试使用）
# ═══════════════════════════════════════════════════════════════


@pytest.fixture
def sample_sections() -> dict[str, str]:
    """一份包含 5 大必备章节的示例文档内容"""
    return {
        "招标公告": (
            "本采购项目采用公开招标方式，欢迎符合资格条件的供应商参加投标。"
            "本次采购预算金额为500万元。"
        ),
        "招标范围": (
            "本次采购内容包括XXX系统建设及运维服务。"
            "项目背景：为提升信息化水平，需建设一套全新的管理系统。"
        ),
        "资格要求": (
            "1.投标人应具有独立承担民事责任的能力。\n"
            "2.投标人必须为本市注册企业，注册资本不低于1000万元。\n"
            "3.本项目不接受联合体投标。\n"
            "4.投标人须具有ISO9001质量管理体系认证。"
        ),
        "评审办法": (
            "本项目采用综合评分法。\n"
            "评分标准：技术方案40分，价格30分，业绩30分。\n"
            "指定品牌XXXX作为评分参考标准。"
        ),
        "投标须知": (
            "投标截止时间：2026年7月1日9:00。\n"
            "投标有效期：自投标截止日起90天。\n"
            "投标保证金：人民币10万元。"
        ),
    }


@pytest.fixture
def sample_docx_path() -> str:
    """创建一份真实的 Word 测试文档，返回路径"""
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_paragraph("本采购项目采用公开招标方式，欢迎合格供应商投标。")
    doc.add_heading("第二章 招标范围", level=1)
    doc.add_paragraph("采购内容详见附件。")
    doc.add_heading("第三章 投标人资格要求", level=1)
    doc.add_paragraph("1. 独立承担民事责任的能力。")
    doc.add_paragraph("2. 本市注册企业优先。")
    doc.add_heading("第四章 评审办法", level=1)
    doc.add_paragraph("综合评分法。指定品牌XXXX产品。")
    doc.add_heading("第五章 投标人须知", level=1)
    doc.add_paragraph("投标截止时间：2026年7月1日。")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    doc.save(tmp.name)

    yield tmp.name

    Path(tmp.name).unlink(missing_ok=True)
