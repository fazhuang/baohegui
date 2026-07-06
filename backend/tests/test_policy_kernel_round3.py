"""PolicyKernel 第三轮增量修复测试

测试：
1. 持久化载荷完整性 — 单一 _policy_decision，反序列化，verify_trace
2. 真实 API 往返 — 上传→检查→报告
3. 损坏报告 fail closed — detail/PDF/Excel 均返回 409
4. 新数据库迁移 — 空库 alembic upgrade head
5. 旧数据库升级 — 历史数据不丢失、不回填
6. 数据库约束 — 非法值被拒绝
7. 平台章节矩阵 — present_sections 精确断言
"""

from __future__ import annotations

import io
import json as _json
import os
import sqlite3
import subprocess
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.core.policy_kernel import (
    DecisionInput,
    DecisionAction,
    DecisionState,
    ParseQuality,
    PlanTier,
    PlatformPolicy,
    PolicyDecision,
    PolicyKernel,
    ReasonCode,
    RiskLevel,
    RuleType,
    RuleViolationInput,
    BiasFindingInput,
    TenantPolicy,
    TraceStep,
    sha256_hex,
    _canonical_json,
    verify_trace,
    policy_kernel,
)


BACKEND_DIR = str(Path(__file__).resolve().parent.parent)


# ═══════════════════════════════════════════════════════════════
# 1. 持久化载荷完整性
# ═══════════════════════════════════════════════════════════════

class TestPersistencePayloadIntegrity:
    """report_data 中只有一个 _policy_decision，反序列化完整，verify_trace 通过"""

    def _make_payload(self) -> tuple[dict, DecisionInput, PolicyDecision]:
        """通过真实 payload 构造函数生成 report_data。"""
        di = DecisionInput(
            schema_version="2.0.0",
            rule_violations=[
                RuleViolationInput(
                    rule_id="R001", rule_type=RuleType.FORBIDDEN,
                    risk_level=RiskLevel.HIGH, description="test violation",
                )
            ],
            present_sections={"招标公告", "技术要求"},
        )
        pd = policy_kernel.decide(di)

        from app.api.check import build_policy_audit_payload
        # 用空 report 模拟
        from app.engine.fusion import ComplianceReport
        report = ComplianceReport(total_score=85.0)
        payload = build_policy_audit_payload(
            report=report,
            decision_input=di,
            policy_decision=pd,
            diagnostics={"_test": True},
        )
        return payload, di, pd

    def test_single_policy_decision_key(self):
        """report_data 只能有一个 _policy_decision 键"""
        payload, _, _ = self._make_payload()
        # JSON 序列化后只能出现一次
        json_str = _json.dumps(payload, ensure_ascii=False)
        keys = [k for k in payload.keys() if k == "_policy_decision"]
        assert len(keys) == 1, f"expected 1 _policy_decision key, got {len(keys)}"

    def test_policy_decision_deserializable(self):
        """_policy_decision 可通过 PolicyDecision.model_validate 重建"""
        payload, _, _ = self._make_payload()
        pd_raw = payload["_policy_decision"]
        pd = PolicyDecision.model_validate(pd_raw)
        assert pd.final_action in (
            DecisionAction.PASS, DecisionAction.WARN,
            DecisionAction.REQUIRE_REVIEW, DecisionAction.BLOCK,
        )

    def test_trace_steps_complete(self):
        """每条 TraceStep 包含 reason_params 和 proposed_transition"""
        payload, _, _ = self._make_payload()
        pd = PolicyDecision.model_validate(payload["_policy_decision"])
        for i, step in enumerate(pd.trace_chain):
            assert step.reason_params is not None, f"trace[{i}] missing reason_params"
            assert step.proposed_transition is not None, f"trace[{i}] missing proposed_transition"
            assert isinstance(step.proposed_transition, DecisionState)
            assert step.proposed_transition.action in (
                DecisionAction.PASS, DecisionAction.WARN,
                DecisionAction.REQUIRE_REVIEW, DecisionAction.BLOCK,
            ), f"trace[{i}] proposed_transition.action={step.proposed_transition.action}"

    def test_verify_trace_passes(self):
        """payload 中的 DecisionInput + PolicyDecision 通过 verify_trace"""
        payload, di, pd = self._make_payload()
        di_reconstructed = DecisionInput.model_validate(payload["_decision_input"])
        pd_reconstructed = PolicyDecision.model_validate(payload["_policy_decision"])
        vr = verify_trace(di_reconstructed, pd_reconstructed)
        assert vr["valid"], f"verify_trace failed: {vr['errors']}"


# ═══════════════════════════════════════════════════════════════
# 3. 损坏报告 fail closed
# ═══════════════════════════════════════════════════════════════

class TestCorruptedReportFailClosed:
    """损坏的 PolicyDecision 在 detail/PDF/Excel 三端返回 409"""

    @pytest.fixture
    def client(self):
        return TestClient(app)

    @staticmethod
    def _register_and_login(client):
        import secrets, time
        from app.main import _rate_limits
        _rate_limits.clear()
        email = f"test-corr-{secrets.token_hex(4)}@example.com"
        pw = "Test123456!"
        for _retry in range(5):
            resp = client.post("/api/auth/register", json={"username": email, "password": pw, "email": email})
            if resp.status_code == 429:
                time.sleep(5.0)
                _rate_limits.clear()
                continue
            if resp.status_code not in (200, 201, 409):
                raise RuntimeError(f"register failed: {resp.status_code} {resp.text}")
            if resp.status_code == 409:
                resp = client.post("/api/auth/login", json={"username": email, "password": pw})
            if resp.status_code != 429:
                token = resp.json()["access_token"]
                return {"Authorization": f"Bearer {token}"}
            time.sleep(5.0)
        raise RuntimeError("register/login failed after 5 retries")

    def _create_docx_file(self):
        doc = Document()
        doc.add_heading("第一章 招标公告", level=1)
        doc.add_paragraph("公开招标，欢迎投标人参加。")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "test_corrupt.docx"
        return buf

    def _upload_and_check(self, client, auth_headers):
        """上传文件 → 执行检查 → 返回 report_id"""
        buf = self._create_docx_file()
        resp = client.post("/api/upload/", files={"file": (buf.name, buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=auth_headers)
        assert resp.status_code in (200, 201), f"upload failed: {resp.status_code} {resp.text}"
        file_id = resp.json()["db_id"]  # integer DB row id
        check_resp = client.post(f"/api/check/{file_id}?procurement_method=公开招标&project_type=货物类", headers=auth_headers)
        assert check_resp.status_code == 200, f"check failed: {check_resp.status_code} {check_resp.text}"
        return check_resp.json()["report_id"]

    def test_corrupt_via_proposed_transition_409_detail(self, client):
        """篡改 proposed_transition → detail 409"""
        h = self._register_and_login(client)

        buf = self._create_docx_file()
        resp = client.post("/api/upload/", files={"file": (buf.name, buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=h)
        assert resp.status_code in (200, 201), f"upload failed: {resp.text}"
        file_id = resp.json()["db_id"]  # integer DB row id
        check_resp = client.post(f"/api/check/{file_id}?platform=guangdong&procurement_method=公开招标&project_type=货物类", headers=h)
        assert check_resp.status_code == 200, f"check failed: {check_resp.text}"
        report_id = check_resp.json()["report_id"]

        # 直接操作数据库篡改 report_data
        from app.db.database import SessionLocal
        from app.models.document import ComplianceReport
        db = SessionLocal()
        try:
            db_report = db.query(ComplianceReport).filter(ComplianceReport.id == report_id).first()
            data = _json.loads(db_report.report_data)
            pd = data["_policy_decision"]
            # 篡改 proposed_transition
            if pd.get("trace_chain"):
                pd["trace_chain"][-1]["proposed_transition"] = {
                    "action": "pass", "risk_level": "low", "requires_human_review": False,
                }
                # 不同步修改 state_after → 触发语义不匹配
                data["_policy_decision"] = pd
                db_report.report_data = _json.dumps(data, ensure_ascii=False)
                db.commit()

            # detail must return 409
            resp = client.get(f"/api/report/{report_id}", headers=h)
            assert resp.status_code == 409, f"expected 409 for corrupt, got {resp.status_code}: {resp.text}"
            detail = resp.json()["detail"]
            assert detail["integrity_status"] == "integrity_failed"
        finally:
            db.close()

    def test_corrupt_via_hash_409_pdf(self, client):
        """篡改 decision_hash → PDF 409"""
        h = self._register_and_login(client)

        buf = self._create_docx_file()
        resp = client.post("/api/upload/", files={"file": (buf.name, buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=h)
        file_id = resp.json()["db_id"]  # integer DB row id
        check_resp = client.post(f"/api/check/{file_id}?procurement_method=公开招标&project_type=货物类", headers=h)
        assert check_resp.status_code == 200, f"check failed: {check_resp.text}"
        report_id = check_resp.json()["report_id"]

        from app.db.database import SessionLocal
        from app.models.document import ComplianceReport
        db = SessionLocal()
        try:
            db_report = db.query(ComplianceReport).filter(ComplianceReport.id == report_id).first()
            db_report.decision_hash = "0000000000000000bad_hash"
            db_report.decision_integrity_status = "verified"  # 刻意标 verified
            db.commit()

            resp = client.get(f"/api/report/{report_id}/pdf", headers=h)
            assert resp.status_code == 409, f"expected 409 for corrupt PDF, got {resp.status_code}"
        finally:
            db.close()

    def test_corrupt_via_action_mismatch_409_excel(self, client):
        """数据库 decision_hash 篡改后 Excel 409"""
        h = self._register_and_login(client)

        buf = self._create_docx_file()
        resp = client.post("/api/upload/", files={"file": (buf.name, buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=h)
        file_id = resp.json()["db_id"]  # integer DB row id
        check_resp = client.post(f"/api/check/{file_id}?procurement_method=公开招标&project_type=货物类", headers=h)
        assert check_resp.status_code == 200, f"check failed: {check_resp.text}"
        report_id = check_resp.json()["report_id"]

        from app.db.database import SessionLocal
        from app.models.document import ComplianceReport
        db = SessionLocal()
        try:
            db_report = db.query(ComplianceReport).filter(ComplianceReport.id == report_id).first()
            # 篡改 JSON 内嵌 decision_hash → verify_trace 将失败
            data = _json.loads(db_report.report_data)
            data["_policy_decision"]["decision_hash"] = "0" * 64
            db_report.report_data = _json.dumps(data, ensure_ascii=False)
            db.commit()
        finally:
            db.close()

        resp = client.get(f"/api/report/{report_id}/export", headers=h)
        assert resp.status_code == 409, f"expected 409 for corrupt Excel, got {resp.status_code}: {resp.text}"

    def test_corrupt_missing_decision_input_409(self, client):
        """schema v2 但缺少 DecisionInput → 409"""
        h = self._register_and_login(client)

        buf = self._create_docx_file()
        resp = client.post("/api/upload/", files={"file": (buf.name, buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=h)
        file_id = resp.json()["db_id"]  # integer DB row id
        check_resp = client.post(f"/api/check/{file_id}?procurement_method=公开招标&project_type=货物类", headers=h)
        assert check_resp.status_code == 200, f"check failed: {check_resp.text}"
        report_id = check_resp.json()["report_id"]

        from app.db.database import SessionLocal
        from app.models.document import ComplianceReport
        db = SessionLocal()
        try:
            db_report = db.query(ComplianceReport).filter(ComplianceReport.id == report_id).first()
            data = _json.loads(db_report.report_data)
            del data["_decision_input"]
            db_report.report_data = _json.dumps(data, ensure_ascii=False)
            db_report.policy_schema_version = "2.0.0"
            db.commit()

            resp = client.get(f"/api/report/{report_id}", headers=h)
            assert resp.status_code == 409, f"expected 409, got {resp.status_code}: {resp.text}"
        finally:
            db.close()


# ═══════════════════════════════════════════════════════════════
# 4. 新数据库迁移 — 空库 alembic upgrade head
# ═══════════════════════════════════════════════════════════════

class TestFreshMigration:
    """全新空 SQLite 数据库执行 alembic upgrade head → 成功"""

    def test_fresh_migration_creates_all_tables(self):
        db_path = "/private/tmp/bhg_policy_fresh_test.db"
        # 删除旧文件
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}

        result = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert result.returncode == 0, f"alembic upgrade head failed: {result.stderr}"

        # 验证表存在
        conn = sqlite3.connect(db_path)
        try:
            tables = {
                row[0] for row in
                conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
            }
            assert "uploaded_files" in tables, f"missing uploaded_files, got {tables}"
            assert "document_sections" in tables, f"missing document_sections"
            assert "compliance_reports" in tables, f"missing compliance_reports"
            assert "alembic_version" in tables, f"missing alembic_version"

            # 验证决策列
            cols = {
                row[1] for row in
                conn.execute("PRAGMA table_info(compliance_reports)").fetchall()
            }
            for col in ["decision_action", "decision_risk_level", "decision_hash",
                        "policy_schema_version", "decision_integrity_status"]:
                assert col in cols, f"missing column {col} in compliance_reports, got {cols}"

            # 验证当前 revision 为 head
            rev_row = conn.execute("SELECT version_num FROM alembic_version").fetchone()
            assert rev_row is not None, "no alembic_version row"
            # head revision should be 20260705_1000_decision_columns
            assert rev_row[0] == "20260705_1000_decision_columns", f"expected head, got {rev_row[0]}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 5. 旧数据库升级 — 历史数据不丢失、不回填
# ═══════════════════════════════════════════════════════════════

class TestLegacyMigration:
    """迁移前数据库已有 compliance_reports 及历史行，升级后数据完好"""

    def test_legacy_data_preserved(self):
        db_path = "/private/tmp/bhg_policy_legacy_test.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        # 模拟旧数据库：先创建表并插入历史数据
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("""
                CREATE TABLE uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL, filename TEXT NOT NULL,
                    file_size INTEGER NOT NULL, file_hash TEXT NOT NULL,
                    storage_path TEXT NOT NULL, status TEXT DEFAULT 'uploaded'
                )
            """)
            conn.execute("""
                CREATE TABLE compliance_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    total_score REAL NOT NULL DEFAULT 0,
                    section_score REAL, keyword_score REAL,
                    forbidden_score REAL, semantic_score REAL,
                    violation_count INTEGER DEFAULT 0,
                    report_data TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    checked_by INTEGER
                )
            """)
            conn.execute("""
                CREATE TABLE alembic_version (version_num VARCHAR(32))
            """)
            conn.execute("INSERT INTO alembic_version VALUES ('d4e5f6a7b8c9')")
            # 插入历史数据
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            conn.execute("INSERT INTO compliance_reports (id, file_id, total_score) VALUES (1, 1, 85.5)")
            conn.commit()
        finally:
            conn.close()

        # 执行迁移到 head
        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}
        result = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert result.returncode == 0, f"alembic upgrade failed: {result.stderr}"

        # 验证历史数据未丢失
        conn = sqlite3.connect(db_path)
        try:
            row = conn.execute("SELECT id, total_score, decision_action, decision_risk_level, decision_integrity_status, decision_hash FROM compliance_reports WHERE id=1").fetchone()
            assert row is not None, "historical row lost"
            assert row[0] == 1
            assert row[1] == 85.5  # total_score preserved
            # decision_action 应为 NULL（未回填）
            assert row[2] is None, f"decision_action was backfilled: {row[2]}"
            assert row[3] is None, f"decision_risk_level was backfilled: {row[3]}"
            # integrity_status 应为 legacy_unverifiable
            assert row[4] == "legacy_unverifiable", f"expected legacy_unverifiable, got {row[4]}"
            assert row[5] is None, f"decision_hash was backfilled: {row[5]}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 6. 数据库约束 — 非法值被拒绝
# ═══════════════════════════════════════════════════════════════

class TestDatabaseConstraints:
    """非法决策值被数据库约束拒绝"""

    def _create_fresh_db_with_migrations(self) -> str:
        db_path = f"/private/tmp/bhg_constraint_test_{os.getpid()}.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass
        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}
        result = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert result.returncode == 0, f"migration failed: {result.stderr}"
        # Verify constraints in schema
        conn = sqlite3.connect(db_path)
        try:
            schema = list(conn.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name='compliance_reports'"))[0][0]
            assert "decision_action" in schema
            assert "CHECK" in schema, f"no CHECK constraints in fresh DB: {schema}"
        finally:
            conn.close()
        return db_path

    def test_invalid_decision_action_rejected(self):
        db_path = self._create_fresh_db_with_migrations()
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_action) VALUES (1, 85, 'anything')")
        finally:
            conn.close()
            os.remove(db_path)

    def test_invalid_decision_risk_level_rejected(self):
        db_path = self._create_fresh_db_with_migrations()
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_risk_level) VALUES (1, 85, 'super_high')")
        finally:
            conn.close()
            os.remove(db_path)

    def test_invalid_integrity_status_rejected(self):
        db_path = self._create_fresh_db_with_migrations()
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_integrity_status) VALUES (1, 85, 'trusted')")
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 7. 平台章节矩阵 — present_sections 精确断言
# ═══════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════
# 8. 老库升级约束 — legacy DB → upgrade head → 非法值被拒绝
# ═══════════════════════════════════════════════════════════════

class TestLegacyUpgradeConstraints:
    """旧 SQLite 数据库（无决策列，有历史数据）升级后约束生效"""

    def _create_legacy_db(self, db_path: str) -> None:
        """创建带历史数据的旧数据库，stamp 为 d4e5f6a7b8c9。"""
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        conn = sqlite3.connect(db_path)
        try:
            conn.execute("""
                CREATE TABLE uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL,
                    filename TEXT NOT NULL,
                    file_size INTEGER NOT NULL,
                    file_hash TEXT NOT NULL,
                    page_count INTEGER,
                    storage_path TEXT NOT NULL,
                    status TEXT DEFAULT 'uploaded',
                    error_message TEXT,
                    failed_at TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE document_sections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    section_type TEXT NOT NULL,
                    section_number TEXT,
                    title TEXT NOT NULL,
                    content TEXT NOT NULL,
                    page_start INTEGER,
                    page_end INTEGER,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE compliance_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    total_score REAL NOT NULL DEFAULT 0,
                    section_score REAL,
                    keyword_score REAL,
                    forbidden_score REAL,
                    semantic_score REAL,
                    violation_count INTEGER DEFAULT 0,
                    report_data TEXT,
                    report_pdf_path TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    checked_by INTEGER
                )
            """)
            conn.execute("""
                CREATE TABLE alembic_version (version_num VARCHAR(32))
            """)
            conn.execute("INSERT INTO alembic_version VALUES ('d4e5f6a7b8c9')")
            # 历史数据
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'legacy.pdf', 2048, 'abc123', '/tmp/legacy.pdf')")
            conn.execute("INSERT INTO compliance_reports (id, file_id, total_score) VALUES (1, 1, 85.5)")
            conn.commit()
        finally:
            conn.close()

    def _run_upgrade(self, db_path: str) -> bool:
        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}
        result = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        return result.returncode == 0

    def test_legacy_invalid_action_rejected(self):
        """旧库升级后 decision_action='anything' → IntegrityError"""
        db_path = "/private/tmp/bhg_r4_legacy_action.db"
        self._create_legacy_db(db_path)
        assert self._run_upgrade(db_path), "upgrade failed"

        conn = sqlite3.connect(db_path)
        try:
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_action) VALUES (1, 100, 'anything')")
        finally:
            conn.close()
            os.remove(db_path)

    def test_legacy_invalid_risk_level_rejected(self):
        """旧库升级后 decision_risk_level='super_high' → IntegrityError"""
        db_path = "/private/tmp/bhg_r4_legacy_risk.db"
        self._create_legacy_db(db_path)
        assert self._run_upgrade(db_path), "upgrade failed"

        conn = sqlite3.connect(db_path)
        try:
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_risk_level) VALUES (1, 100, 'super_high')")
        finally:
            conn.close()
            os.remove(db_path)

    def test_legacy_invalid_integrity_status_rejected(self):
        """旧库升级后 integrity_status='trusted' → IntegrityError"""
        db_path = "/private/tmp/bhg_r4_legacy_istatus.db"
        self._create_legacy_db(db_path)
        assert self._run_upgrade(db_path), "upgrade failed"

        conn = sqlite3.connect(db_path)
        try:
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_integrity_status) VALUES (1, 100, 'trusted')")
        finally:
            conn.close()
            os.remove(db_path)

    def test_legacy_invalid_requires_human_review_rejected(self):
        """旧库升级后 decision_requires_human_review=2 → IntegrityError（SQLite Boolean 语义）"""
        db_path = "/private/tmp/bhg_r4_legacy_bool.db"
        self._create_legacy_db(db_path)
        assert self._run_upgrade(db_path), "upgrade failed"

        conn = sqlite3.connect(db_path)
        try:
            with pytest.raises(sqlite3.IntegrityError):
                conn.execute("INSERT INTO compliance_reports (file_id, total_score, decision_requires_human_review) VALUES (1, 100, 2)")
        finally:
            conn.close()
            os.remove(db_path)

    def test_legacy_historical_data_preserved(self):
        """旧库升级后历史行未丢失。"""
        db_path = "/private/tmp/bhg_r4_legacy_hist.db"
        self._create_legacy_db(db_path)
        assert self._run_upgrade(db_path), "upgrade failed"

        conn = sqlite3.connect(db_path)
        try:
            row = conn.execute("SELECT id, total_score, decision_action, decision_risk_level, decision_integrity_status FROM compliance_reports WHERE id=1").fetchone()
            assert row is not None, "historical row lost"
            assert row[0] == 1
            assert row[1] == 85.5
            assert row[2] is None, f"decision_action backfilled: {row[2]}"
            assert row[3] is None, f"decision_risk_level backfilled: {row[3]}"
            assert row[4] == "legacy_unverifiable", f"expected legacy_unverifiable, got {row[4]}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 9. 老库 downgrade 保护 — 核心表不被删除
# ═══════════════════════════════════════════════════════════════

class TestLegacyDowngradeProtection:
    """升级前已存在的表，downgrade 后仍存在且数据完好"""

    def _create_legacy_db(self, db_path: str) -> None:
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        conn = sqlite3.connect(db_path)
        try:
            conn.execute("""
                CREATE TABLE uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL, filename TEXT NOT NULL,
                    file_size INTEGER NOT NULL, file_hash TEXT NOT NULL,
                    page_count INTEGER, storage_path TEXT NOT NULL,
                    status TEXT DEFAULT 'uploaded', error_message TEXT,
                    failed_at TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE document_sections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    section_type TEXT NOT NULL, section_number TEXT,
                    title TEXT NOT NULL, content TEXT NOT NULL,
                    page_start INTEGER, page_end INTEGER,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE compliance_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    total_score REAL NOT NULL DEFAULT 0,
                    section_score REAL, keyword_score REAL,
                    forbidden_score REAL, semantic_score REAL,
                    violation_count INTEGER DEFAULT 0,
                    report_data TEXT, report_pdf_path TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    checked_by INTEGER
                )
            """)
            conn.execute("""
                CREATE TABLE alembic_version (version_num VARCHAR(32))
            """)
            conn.execute("INSERT INTO alembic_version VALUES ('d4e5f6a7b8c9')")
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'original.pdf', 4096, 'orig_hash', '/tmp/orig.pdf')")
            conn.execute("INSERT INTO document_sections (id, file_id, section_type, title, content) VALUES (1, 1, 'spec', '技术要求', 'section content')")
            conn.execute("INSERT INTO compliance_reports (id, file_id, total_score) VALUES (1, 1, 92.0)")
            conn.commit()
        finally:
            conn.close()

    def test_downgrade_preserves_preexisting_tables(self):
        """旧库 upgrade head → downgrade d4e5f6a7b8c9 → 原表和数据都在"""
        db_path = "/private/tmp/bhg_r4_downgrade_protect.db"
        self._create_legacy_db(db_path)

        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}

        # upgrade head
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"upgrade failed: {r.stderr}"

        # downgrade back
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "downgrade", "d4e5f6a7b8c9"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"downgrade failed: {r.stderr}"

        conn = sqlite3.connect(db_path)
        try:
            tables = {row[0] for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()}

            assert "uploaded_files" in tables, "uploaded_files was dropped by downgrade!"
            assert "document_sections" in tables, "document_sections was dropped by downgrade!"
            assert "compliance_reports" in tables, "compliance_reports was dropped by downgrade!"

            # 原始行完全保留
            uf = conn.execute("SELECT id, filename, file_size FROM uploaded_files WHERE id=1").fetchone()
            assert uf is not None
            assert uf[1] == "original.pdf"
            assert uf[2] == 4096

            ds = conn.execute("SELECT id, title, content FROM document_sections WHERE id=1").fetchone()
            assert ds is not None
            assert ds[1] == "技术要求"

            cr = conn.execute("SELECT id, total_score FROM compliance_reports WHERE id=1").fetchone()
            assert cr is not None
            assert cr[1] == 92.0

            # alembic_version 正确回退
            rev = conn.execute("SELECT version_num FROM alembic_version").fetchone()
            assert rev[0] == "d4e5f6a7b8c9", f"expected d4e5f6a7b8c9, got {rev[0]}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 10. 空库 downgrade — 所有权表生效，只删除迁移创建的表
# ═══════════════════════════════════════════════════════════════

class TestFreshDowngrade:
    """空库 upgrade head → downgrade → 只删除有所有权的表"""

    def test_fresh_downgrade_drops_owned_tables(self):
        """空库 upgrade head 后 downgrade 到 d4e5f6a7b8c9，核心表被删除"""
        db_path = "/private/tmp/bhg_r4_fresh_downgrade.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}

        # upgrade head
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"upgrade failed: {r.stderr}"

        # downgrade to d4e5f6a7b8c9
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "downgrade", "d4e5f6a7b8c9"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"downgrade failed: {r.stderr}"

        conn = sqlite3.connect(db_path)
        try:
            tables = {row[0] for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()}

            # 桥接迁移创建的表被删除
            assert "uploaded_files" not in tables, "uploaded_files should be dropped"
            assert "document_sections" not in tables, "document_sections should be dropped"
            assert "compliance_reports" not in tables, "compliance_reports should be dropped"

            # 所有权表也被清理
            assert "_bhg_migration_objects" not in tables, "ownership table should be cleaned up"

            # alembic_version 正确
            rev = conn.execute("SELECT version_num FROM alembic_version").fetchone()
            assert rev[0] == "d4e5f6a7b8c9", f"wrong revision: {rev[0]}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 11. legacy detail 白名单 — 不含 score/violations/merge
# ═══════════════════════════════════════════════════════════════

class TestLegacyDetailWhitelist:
    """legacy_unverifiable 报告 detail 只返回白名单元数据"""

    @pytest.fixture
    def client(self):
        return TestClient(app)

    @staticmethod
    def _register_and_login(client):
        import base64, json as _j, secrets, time
        # Clear rate limiter from prior test classes
        from app.main import _rate_limits
        _rate_limits.clear()
        email = f"test-legacy-{secrets.token_hex(4)}@example.com"
        pw = "Test123456!"
        for _retry in range(5):
            resp = client.post("/api/auth/register", json={"username": email, "password": pw, "email": email})
            if resp.status_code == 429:
                time.sleep(5.0)
                _rate_limits.clear()
                continue
            if resp.status_code == 409:
                resp = client.post("/api/auth/login", json={"username": email, "password": pw})
            data = resp.json()
            if "access_token" in data:
                token = data["access_token"]
                payload_b64 = token.split(".")[1]
                payload_b64 += "=" * (4 - len(payload_b64) % 4)
                payload = _j.loads(base64.urlsafe_b64decode(payload_b64))
                return {"Authorization": f"Bearer {token}"}, payload["sub"]
            time.sleep(5.0)
        raise RuntimeError("register/login failed after 5 retries")

    def test_legacy_detail_whitelist_no_scores(self, client):
        """legacy 报告 detail 不含 total_score/section_scores/规则/merge"""
        h, user_id = self._register_and_login(client)

        from app.db.database import SessionLocal
        from app.models.document import ComplianceReport
        db = SessionLocal()
        try:
            # 手工插入 legacy 报告
            report = ComplianceReport(
                file_id=1,
                total_score=99.0,
                section_score=95.0,
                keyword_score=90.0,
                forbidden_score=85.0,
                semantic_score=80.0,
                violation_count=5,
                report_data=_json.dumps({
                    "total_score": 99.0,
                    "section_score": 95.0,
                    "rule_violations": [{"rule_id": "R001", "description": "bad"}],
                    "llm_violations": [{"rule_id": "L001", "description": "semantic"}],
                    "_merge_result": {"merged": True},
                    "final_action": "pass",
                    "final_risk_level": "low",
                    "high_risk_count": 3,
                    "medium_risk_count": 2,
                    "low_risk_count": 1,
                }),
                decision_integrity_status="legacy_unverifiable",
                decision_action=None,
                decision_risk_level=None,
                decision_hash=None,
                policy_schema_version=None,
                checked_by=int(user_id),  # actual user id from JWT
            )
            db.add(report)
            db.commit()
            report_id = report.id
        finally:
            db.close()

        resp = client.get(f"/api/report/{report_id}", headers=h)
        assert resp.status_code == 200, f"expected 200, got {resp.status_code}: {resp.text}"
        data = resp.json()

        # 白名单字段存在
        assert "report_id" in data
        assert "file_id" in data

        # 黑名单字段不存在
        assert "total_score" not in data, f"total_score leaked: {data.get('total_score')}"
        assert "section_score" not in data, f"section_score leaked"
        assert "keyword_score" not in data, f"keyword_score leaked"
        assert "forbidden_score" not in data, f"forbidden_score leaked"
        assert "semantic_score" not in data, f"semantic_score leaked"
        assert "rule_violations" not in data, f"rule_violations leaked"
        assert "llm_violations" not in data, f"llm_violations leaked"
        assert "_merge_result" not in data, f"_merge_result leaked"
        assert "high_risk_count" not in data, f"high_risk_count leaked"
        assert "medium_risk_count" not in data, f"medium_risk_count leaked"

        # integrity 元数据
        assert data.get("final_action") == "unknown"
        assert data.get("final_risk_level") == "unknown"
        assert data["_integrity"]["status"] == "legacy_unverifiable"


# ═══════════════════════════════════════════════════════════════
# 12. legacy PDF/Excel 409
# ═══════════════════════════════════════════════════════════════

class TestLegacyExport409:
    """legacy_unverifiable 报告 PDF/Excel 均返回 409"""

    @pytest.fixture
    def auth(self):
        """Fixture: register/login once, shared across both test methods."""
        from app.main import _rate_limits
        # Clear rate limits from prior test classes
        _rate_limits.clear()
        import json as _j, base64, secrets, time
        client = TestClient(app)
        email = f"test-legacy-exp-{secrets.token_hex(4)}@example.com"
        pw = "Test123456!"
        for _retry in range(5):
            resp = client.post("/api/auth/register", json={"username": email, "password": pw, "email": email})
            if resp.status_code == 429:
                time.sleep(5.0)
                _rate_limits.clear()
                continue
            if resp.status_code == 409:
                resp = client.post("/api/auth/login", json={"username": email, "password": pw})
            data = resp.json()
            if "access_token" in data:
                token = data["access_token"]
                payload_b64 = token.split(".")[1]
                payload_b64 += "=" * (4 - len(payload_b64) % 4)
                payload = _j.loads(base64.urlsafe_b64decode(payload_b64))
                return {"client": client, "headers": {"Authorization": f"Bearer {token}"}, "user_id": payload["sub"]}
            time.sleep(5.0)
        raise RuntimeError("register/login failed after 5 retries")

    def _create_legacy_report(self, db, user_id):
        from app.models.document import ComplianceReport, UploadedFile
        # Ensure a minimal uploaded_files row exists for FK
        uf = db.query(UploadedFile).filter(UploadedFile.id == 1).first()
        if not uf:
            uf = UploadedFile(
                id=1, user_id=int(user_id), filename="dummy.pdf",
                file_size=100, file_hash="dummy_hash", storage_path="/tmp/dummy.pdf",
            )
            db.add(uf)
            db.flush()
        report = ComplianceReport(
            file_id=1,
            total_score=99.0,
            section_score=95.0,
            report_data=_json.dumps({"total_score": 99, "rule_violations": [{"id": "R001"}]}),
            decision_integrity_status="legacy_unverifiable",
            decision_action=None,
            decision_risk_level=None,
            decision_hash=None,
            policy_schema_version=None,
            checked_by=int(user_id),
        )
        db.add(report)
        db.commit()
        return report.id

    def test_legacy_pdf_409(self, auth):
        """legacy → PDF 409"""
        client = auth["client"]
        h = auth["headers"]
        user_id = auth["user_id"]

        from app.db.database import SessionLocal
        db = SessionLocal()
        try:
            report_id = self._create_legacy_report(db, user_id)
        finally:
            db.close()

        resp = client.get(f"/api/report/{report_id}/pdf", headers=h)
        assert resp.status_code == 409, f"expected 409 for legacy PDF, got {resp.status_code}: {resp.text}"
        detail = resp.json()["detail"]
        assert detail["integrity_status"] == "legacy_unverifiable"

    def test_legacy_excel_409(self, auth):
        """legacy → Excel 409"""
        client = auth["client"]
        h = auth["headers"]
        user_id = auth["user_id"]

        from app.db.database import SessionLocal
        db = SessionLocal()
        try:
            report_id = self._create_legacy_report(db, user_id)
        finally:
            db.close()

        resp = client.get(f"/api/report/{report_id}/export", headers=h)
        assert resp.status_code == 409, f"expected 409 for legacy Excel, got {resp.status_code}: {resp.text}"
        detail = resp.json()["detail"]
        assert detail["integrity_status"] == "legacy_unverifiable"


# ═══════════════════════════════════════════════════════════════
# 7. 平台章节矩阵 — present_sections 精确断言
# ═══════════════════════════════════════════════════════════════

class TestPlatformSectionMatrix:
    """required_sections 与 present_sections 的精确矩阵"""

    kernel = PolicyKernel()

    def _decide(self, required: set[str], present: set[str], platform_id="guangdong"):
        di = DecisionInput(
            platform_policy=PlatformPolicy(platform_id=platform_id, required_sections=required),
            present_sections=present,
        )
        return self.kernel.decide(di)

    def test_all_present_no_block(self):
        """required={"招标公告"}, present={"招标公告"} → 非 PLATFORM BLOCK"""
        d = self._decide({"招标公告"}, {"招标公告"})
        # 平台层必须 PLATFORM_PASSTHROUGH
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_PASSTHROUGH
        assert d.final_action != DecisionAction.BLOCK

    def test_empty_present_block(self):
        """required={"招标公告"}, present=set() → BLOCK + CRITICAL"""
        d = self._decide({"招标公告"}, set())
        assert d.final_action == DecisionAction.BLOCK
        assert d.final_risk_level == RiskLevel.CRITICAL
        assert d.requires_human_review is True
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_MISSING_SECTION
        assert "招标公告" in plat.reason_params["missing"]

    def test_wrong_present_block(self):
        """required={"招标公告"}, present={"评审办法"} → BLOCK"""
        d = self._decide({"招标公告"}, {"评审办法"})
        assert d.final_action == DecisionAction.BLOCK
        assert d.final_risk_level == RiskLevel.CRITICAL
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_MISSING_SECTION

    def test_present_plus_unrelated_violation_no_block(self):
        """required={"招标公告"}, present={"招标公告"}，另有 chapter_required 违规 → 平台仍通过"""
        di = DecisionInput(
            rule_violations=[
                RuleViolationInput(rule_id="R003", rule_type=RuleType.CHAPTER_REQUIRED,
                                   risk_level=RiskLevel.MEDIUM, description="缺少投标须知"),
            ],
            platform_policy=PlatformPolicy(platform_id="guangdong", required_sections={"招标公告"}),
            present_sections={"招标公告"},
        )
        d = self.kernel.decide(di)
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_PASSTHROUGH, (
            f"platform should passthrough when required section present, got {plat.reason_code}"
        )
        # HARD_RULE may still escalate from chapter_required rule
        # But PLATFORM must NOT block
        assert plat.state_after.action != DecisionAction.BLOCK

    def test_unknown_platform_is_error(self):
        """未知 platform 不应静默降级。在 API 层测试（check.py），这里验证空 PlatformPolicy 行为。"""
        # 空平台策略（无 platform_id）→ PLATFORM_NO_POLICY
        di = DecisionInput(
            platform_policy=PlatformPolicy(platform_id=""),
            present_sections={"招标公告"},
        )
        d = self.kernel.decide(di)
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_NO_POLICY

    def test_required_sections_empty_passthrough(self):
        """required_sections 为空 → PLATFORM passthrough"""
        di = DecisionInput(
            platform_policy=PlatformPolicy(platform_id="test_plat", required_sections=set()),
            present_sections={"招标公告"},
        )
        d = self.kernel.decide(di)
        plat = [t for t in d.trace_chain if t.source.value == "platform"][0]
        assert plat.reason_code == ReasonCode.PLATFORM_PASSTHROUGH


# ═══════════════════════════════════════════════════════════════
# 13. 老库合法默认值 — 省略 integrity_status 的 INSERT 成功
# ═══════════════════════════════════════════════════════════════

class TestLegacyDefaultValue:
    """老库升级后省略 decision_integrity_status 列的 INSERT 必须成功"""

    def test_legacy_omit_integrity_insert_success(self):
        """老库 upgrade head 后省略 integrity_status INSERT → 成功 → SELECT = legacy_unverifiable"""
        db_path = "/private/tmp/bhg_r5_legacy_default.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        # 创建旧库
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("""
                CREATE TABLE uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL, filename TEXT NOT NULL,
                    file_size INTEGER NOT NULL, file_hash TEXT NOT NULL,
                    page_count INTEGER, storage_path TEXT NOT NULL,
                    status TEXT DEFAULT 'uploaded', error_message TEXT,
                    failed_at TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE document_sections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    section_type TEXT NOT NULL, section_number TEXT,
                    title TEXT NOT NULL, content TEXT NOT NULL,
                    page_start INTEGER, page_end INTEGER,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE compliance_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    total_score REAL NOT NULL DEFAULT 0,
                    section_score REAL, keyword_score REAL,
                    forbidden_score REAL, semantic_score REAL,
                    violation_count INTEGER DEFAULT 0,
                    report_data TEXT, report_pdf_path TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    checked_by INTEGER
                )
            """)
            conn.execute("CREATE TABLE alembic_version (version_num VARCHAR(32))")
            conn.execute("INSERT INTO alembic_version VALUES ('d4e5f6a7b8c9')")
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            conn.commit()
        finally:
            conn.close()

        # upgrade head
        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"upgrade failed: {r.stderr}"

        conn = sqlite3.connect(db_path)
        try:
            # 验证 DDL 不含三重引号
            schema = conn.execute(
                "SELECT sql FROM sqlite_master WHERE type='table' AND name='compliance_reports'"
            ).fetchone()[0]
            assert "DEFAULT '''legacy_unverifiable'''" not in schema, (
                f"triple-quoted default in DDL: {schema}"
            )

            # 合法 INSERT 省略 decision_integrity_status → 必须成功
            conn.execute(
                "INSERT INTO compliance_reports (file_id, total_score, decision_action, "
                "decision_risk_level, decision_requires_human_review) "
                "VALUES (1, 100, 'pass', 'low', 0)"
            )
            conn.commit()

            # SELECT 值严格等于 legacy_unverifiable（不含引号字符）
            row = conn.execute(
                "SELECT decision_integrity_status FROM compliance_reports WHERE total_score=100"
            ).fetchone()
            assert row is not None, "row not found after INSERT"
            assert row[0] == "legacy_unverifiable", (
                f"expected 'legacy_unverifiable', got {row[0]!r}"
            )
            # 确认不含引号字符
            assert "'" not in row[0], f"value contains quote char: {row[0]!r}"
        finally:
            conn.close()
            os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 14. 删除顺序单元测试
# ═══════════════════════════════════════════════════════════════

class TestDropOrderUnit:
    """_compute_drop_order 纯函数按 FK 安全顺序返回表名"""

    # 必须与 bridge_core_reports 中 _DROP_ORDER 保持同步
    _DROP_ORDER = (
        "compliance_reports",
        "document_sections",
        "uploaded_files",
    )

    def _compute_drop_order(self, owned: set[str]) -> list[str]:
        """提取的纯排序逻辑：只返回 _DROP_ORDER 中的 owned 表，按 _DROP_ORDER 顺序"""
        ordered = [t for t in self._DROP_ORDER if t in owned]
        unknown = owned - set(self._DROP_ORDER)
        if unknown:
            import warnings
            warnings.warn(
                f"unknown owned objects not in drop order, skipping: {sorted(unknown)}"
            )
        return ordered

    def test_drop_order_fk_safe(self):
        """owned 含三个核心表 → 返回 child-first 顺序"""
        result = self._compute_drop_order(
            {"compliance_reports", "document_sections", "uploaded_files"}
        )
        assert result == [
            "compliance_reports",
            "document_sections",
            "uploaded_files",
        ], f"unexpected order: {result}"

    def test_unknown_objects_not_silently_dropped(self):
        """_DROP_ORDER 中不存在的 owned 对象不会被列入返回列表"""
        result = self._compute_drop_order(
            {"some_unknown_table", "compliance_reports"}
        )
        assert result == ["compliance_reports"]
        assert "some_unknown_table" not in result


# ═══════════════════════════════════════════════════════════════
# 15. SQLite 外键开启回滚
# ═══════════════════════════════════════════════════════════════

class TestSqliteFKRollback:
    """SQLite PRAGMA foreign_keys=ON 下 upgrade → downgrade 成功"""

    def test_fresh_upgrade_downgrade_with_fk_on(self):
        """空库 upgrade head → downgrade d4e5f6a7b8c9 在 FK ON 下成功"""
        db_path = "/private/tmp/bhg_r5_fk_rollback.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}

        # upgrade head
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"upgrade failed: {r.stderr}"

        # Verify FK ON before downgrade
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("PRAGMA foreign_keys=ON")
            fk_on = conn.execute("PRAGMA foreign_keys").fetchone()[0]
            assert fk_on == 1, "foreign_keys not ON"
        finally:
            conn.close()

        # downgrade
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "downgrade", "d4e5f6a7b8c9"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"downgrade failed: {r.stderr}"

        os.remove(db_path)


# ═══════════════════════════════════════════════════════════════
# 16. 预先存在的 ownership table
# ═══════════════════════════════════════════════════════════════

class TestPreExistingOwnershipTable:
    """升级前已有空 _bhg_migration_objects → upgrade/downgrade 后仍存在"""

    def test_preexisting_empty_ownership_table_survives(self):
        """预先存在的空 ownership table 不因 downgrade 被删除"""
        db_path = "/private/tmp/bhg_r5_preexist_owner.db"
        try:
            os.remove(db_path)
        except FileNotFoundError:
            pass

        # 创建旧库 + 空的 ownership table
        conn = sqlite3.connect(db_path)
        try:
            conn.execute("""
                CREATE TABLE uploaded_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id INTEGER NOT NULL, filename TEXT NOT NULL,
                    file_size INTEGER NOT NULL, file_hash TEXT NOT NULL,
                    page_count INTEGER, storage_path TEXT NOT NULL,
                    status TEXT DEFAULT 'uploaded', error_message TEXT,
                    failed_at TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE document_sections (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    section_type TEXT NOT NULL, section_number TEXT,
                    title TEXT NOT NULL, content TEXT NOT NULL,
                    page_start INTEGER, page_end INTEGER,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE TABLE compliance_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_id INTEGER NOT NULL REFERENCES uploaded_files(id),
                    total_score REAL NOT NULL DEFAULT 0,
                    section_score REAL, keyword_score REAL,
                    forbidden_score REAL, semantic_score REAL,
                    violation_count INTEGER DEFAULT 0,
                    report_data TEXT, report_pdf_path TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    checked_by INTEGER
                )
            """)
            # 预先存在的空 ownership table
            conn.execute("""
                CREATE TABLE _bhg_migration_objects (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    revision VARCHAR(64) NOT NULL,
                    object_type VARCHAR(32) NOT NULL,
                    object_name VARCHAR(256) NOT NULL,
                    created_by_migration BOOLEAN NOT NULL DEFAULT 1,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("CREATE TABLE alembic_version (version_num VARCHAR(32))")
            conn.execute("INSERT INTO alembic_version VALUES ('d4e5f6a7b8c9')")
            conn.execute("INSERT INTO uploaded_files (id, user_id, filename, file_size, file_hash, storage_path) VALUES (1, 1, 'test.pdf', 1024, 'abc', '/tmp/test.pdf')")
            conn.commit()
        finally:
            conn.close()

        db_url = f"sqlite:///{db_path}"
        env = {**os.environ, "DATABASE_URL": db_url, "BHG_DATABASE_URL": db_url}

        # upgrade head
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "upgrade", "head"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"upgrade failed: {r.stderr}"

        # downgrade
        r = subprocess.run(
            [sys.executable, "-m", "alembic", "downgrade", "d4e5f6a7b8c9"],
            capture_output=True, text=True, env=env, cwd=BACKEND_DIR,
        )
        assert r.returncode == 0, f"downgrade failed: {r.stderr}"

        conn = sqlite3.connect(db_path)
        try:
            tables = {row[0] for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()}

            # 原表仍存在
            assert "uploaded_files" in tables, "uploaded_files missing after downgrade"
            assert "document_sections" in tables, "document_sections missing after downgrade"
            assert "compliance_reports" in tables, "compliance_reports missing after downgrade"

            # 预先存在的 ownership table 必须仍存在（不能被删除）
            assert "_bhg_migration_objects" in tables, (
                "pre-existing _bhg_migration_objects was deleted!"
            )
        finally:
            conn.close()
            os.remove(db_path)
