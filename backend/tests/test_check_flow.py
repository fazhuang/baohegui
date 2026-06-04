"""合规检查流程集成测试 — 上传 + 合规审查端到端"""

from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient


def _create_test_docx() -> str:
    """创建一份测试用 Word 招标文件，返回文件路径"""
    doc = Document()
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_paragraph("本采购项目采用公开招标方式，欢迎符合资格条件的供应商参加。")
    doc.add_heading("第二章 招标范围", level=1)
    doc.add_paragraph("采购内容：XX系统建设及运维服务。")
    doc.add_heading("第三章 投标人资格要求", level=1)
    doc.add_paragraph("1. 投标人应具有独立承担民事责任的能力。")
    doc.add_paragraph("2. 本项目接受联合体投标。")
    doc.add_paragraph("3. 投标人须具有ISO9001质量管理体系认证。")
    doc.add_heading("第四章 评审办法", level=1)
    doc.add_paragraph("本项目采用综合评分法。")
    doc.add_paragraph("评分标准：技术方案40分，价格30分，业绩30分。")
    doc.add_heading("第五章 投标人须知", level=1)
    doc.add_paragraph("投标截止时间：2026年7月1日9:00。")
    doc.add_paragraph("投标有效期：自投标截止日起90天。")
    doc.add_paragraph("投标保证金：人民币10万元。")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


class TestUploadFile:
    """文件上传 API 测试"""

    def test_upload_docx_success(self, client: TestClient, auth_headers):
        """上传合法 docx 文件应成功"""
        docx_path = _create_test_docx()
        try:
            with open(docx_path, "rb") as f:
                resp = client.post(
                    "/api/upload/",
                    files={"file": ("test_bid.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                    headers=auth_headers,
                )
            assert resp.status_code == 200, f"Upload failed: {resp.text}"
            data = resp.json()
            assert "file_id" in data
            assert "db_id" in data
            assert data["filename"] == "test_bid.docx"
            assert data["page_count"] is not None
            assert "sections" in data
        finally:
            Path(docx_path).unlink(missing_ok=True)

    def test_upload_without_auth(self, client: TestClient):
        """未认证上传应返回 403/401"""
        docx_path = _create_test_docx()
        try:
            with open(docx_path, "rb") as f:
                resp = client.post(
                    "/api/upload/",
                    files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                )
            # Should be unauthorized — FastAPI bearer scheme returns 403
            assert resp.status_code in (401, 403), f"Expected 401/403, got {resp.status_code}"
        finally:
            Path(docx_path).unlink(missing_ok=True)

    def test_upload_invalid_extension(self, client: TestClient, auth_headers):
        """上传不支持的格式应返回 400"""
        resp = client.post(
            "/api/upload/",
            files={"file": ("test.txt", io.BytesIO(b"not a docx"), "text/plain")},
            headers=auth_headers,
        )
        assert resp.status_code == 400
        assert "不支持的文件格式" in resp.json()["detail"]


class TestCheckFlow:
    """合规检查 API 测试"""

    def _upload_test_file(self, client: TestClient, auth_headers) -> int:
        """Helper: 上传测试文档并返回 db_id"""
        docx_path = _create_test_docx()
        try:
            with open(docx_path, "rb") as f:
                resp = client.post(
                    "/api/upload/",
                    files={"file": ("check_test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                    headers=auth_headers,
                )
            assert resp.status_code == 200, f"Upload failed: {resp.text}"
            return resp.json()["db_id"]
        finally:
            Path(docx_path).unlink(missing_ok=True)

    def test_check_file_returns_score_and_violations(self, client: TestClient, auth_headers):
        """合规检查应返回 total_score 和 violations"""
        file_id = self._upload_test_file(client, auth_headers)

        resp = client.post(f"/api/check/{file_id}", headers=auth_headers)
        assert resp.status_code == 200, f"Check failed: {resp.text}"
        data = resp.json()

        assert "total_score" in data
        assert "total_violations" in data
        assert "high_risk_count" in data
        assert "medium_risk_count" in data
        assert "low_risk_count" in data
        assert "section_score" in data
        assert isinstance(data["total_score"], (int, float))
        assert 0 <= data["total_score"] <= 100

    def test_check_nonexistent_file(self, client: TestClient, auth_headers):
        """检查不存在的文件应返回 404"""
        resp = client.post("/api/check/99999", headers=auth_headers)
        assert resp.status_code == 404

    def test_check_with_industry_params(self, client: TestClient, auth_headers):
        """带行业参数检查应正常工作"""
        file_id = self._upload_test_file(client, auth_headers)

        resp = client.post(
            f"/api/check/{file_id}",
            params={"industries": "it,healthcare", "sector": "政府采购", "procurement_method": "公开招标"},
            headers=auth_headers,
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "total_score" in data
        # 行业参数应出现在返回结果中
        assert data.get("industries") == ["it", "healthcare"]

    def test_check_without_auth(self, client: TestClient, auth_headers):
        """未认证检查文件应返回 401/403"""
        file_id = self._upload_test_file(client, auth_headers)

        resp = client.post(f"/api/check/{file_id}")
        assert resp.status_code in (401, 403)

    def test_check_result_can_be_converted_to_json(self, client: TestClient, auth_headers):
        """检查结果应可被 JSON 序列化"""
        file_id = self._upload_test_file(client, auth_headers)

        resp = client.post(f"/api/check/{file_id}", headers=auth_headers)
        assert resp.status_code == 200, f"Check failed: {resp.text}"

        # 验证 JSON 解析
        import json
        data = json.loads(resp.text)
        assert data["total_score"] is not None
        assert data["total_violations"] is not None
