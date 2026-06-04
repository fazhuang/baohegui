"""招标文件测试文档生成器

用于生成合规/不合规的测试用招标文件（PDF 和 DOCX 格式），
覆盖正常流程、异常流程和边界场景的测试数据需求。
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Optional

import docx
import fitz

# ── 合规文档的内容模板 ──────────────────────────────────────

COMPLIANT_SECTIONS: list[dict] = [
    {
        "title": "第一章 招标公告",
        "content": (
            "本项目采用公开招标方式，欢迎合格的供应商参与投标。\n"
            "采购预算为人民币500万元。\n"
            "本次招标为XXX信息化建设项目。"
        ),
    },
    {
        "title": "第二章 招标范围",
        "content": (
            "本次采购范围为XXX系统建设及配套服务。\n"
            "项目背景：为提升信息化水平，需建设一套全新的管理系统。\n"
            "技术参数要求详见附件。"
        ),
    },
    {
        "title": "第三章 投标人资格要求",
        "content": (
            "1. 投标人应具有独立承担民事责任的能力。\n"
            "2. 投标人应具有良好的商业信誉和健全的财务会计制度。\n"
            "3. 投标人应遵守公平竞争原则，不得以不正当手段参与竞争。\n"
            "4. 本项目接受联合体投标，联合体各方均应满足资格要求。\n"
            "5. 投标人可依法分包，分包比例不得超过30%。\n"
            "6. 中小企业享受价格扣除优惠。"
        ),
    },
    {
        "title": "第四章 评审办法",
        "content": (
            "本项目采用综合评分法。\n"
            "评审标准：技术方案40分，价格35分，商务15分，服务10分。\n"
            "评审标准详见评分细则。\n"
            "节能环保产品给予加分。"
        ),
    },
    {
        "title": "第五章 投标人须知",
        "content": (
            "投标截止时间：2026年7月1日9:00。\n"
            "开标时间：2026年7月1日9:00。\n"
            "投标有效期：自投标截止日起90天。\n"
            "投标保证金：人民币10万元。\n"
            "有下列情形之一的，作废标处理：（一）投标文件逾期送达…\n"
            "若对招标文件有质疑，应在投标截止前10日内提出。\n"
            "踏勘现场时间：2026年6月20日。\n"
            "履约保证金：合同金额的10%。\n"
            "报价有效期90天。\n"
            "投标人应签署廉洁自律承诺书。\n"
            "招标人将通过信用中国网站查询投标人信用记录。"
        ),
    },
    {
        "title": "第六章 合同条款",
        "content": (
            "1. 付款方式：合同签订后支付30%，验收后支付70%。\n"
            "2. 验收程序：项目完成后15个工作日内组织验收。\n"
            "3. 违约责任：任何一方违约应承担合同金额10%的违约金。\n"
            "4. 争议解决方式：双方协商解决，协商不成的提交仲裁委员会仲裁。"
        ),
    },
    {
        "title": "第七章 投标文件格式",
        "content": (
            "投标人应提交以下文件：\n"
            "1. 投标函\n"
            "2. 法定代表人身份证明\n"
            "3. 报价表\n"
            "4. 资格证明文件"
        ),
    },
    {
        "title": "投标保证金",
        "content": "投标保证金金额为人民币10万元。提交方式为银行转账。",
    },
    {
        "title": "报价要求",
        "content": (
            "投标人应对本次采购的全部内容进行报价。\n"
            "报价应包含所有税费、运输费、安装调试费等。"
        ),
    },
    {
        "title": "履约要求",
        "content": (
            "项目实施计划：合同签订后60天内完成。\n"
            "人员配置：项目经理1名，技术工程师2名。"
        ),
    },
    {
        "title": "保密条款",
        "content": (
            "双方应对招标投标过程中知悉的对方商业秘密和保密信息予以保密。\n"
            "保密期限为合同终止后3年。"
        ),
    },
]

# ── 各种违规内容片段 ───────────────────────────────────────

FORBIDDEN_WORDS_CONTENT = (
    "指定品牌XXXX作为唯一授权产品。\n"
    "投标人须为本市注册企业。\n"
    "注册资本不低于1000万元。\n"
    "投标人须获得原厂唯一授权。\n"
    "独家代理权。\n"
    "本省企业优先。"
)

MISSING_KEYWORDS_CONTENT = (
    "欢迎参与本项目。\n"
    "请按要求提交文件。\n"
    "相关标准详见附件。"
)


def _write_docx(sections: list[dict], path: str) -> str:
    """将章节列表写入 DOCX 文件"""
    doc = docx.Document()
    for sec in sections:
        doc.add_heading(sec["title"], level=1)
        for para in sec["content"].split("\n"):
            doc.add_paragraph(para.strip())
    # 确保目录存在
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def _write_pdf(sections: list[dict], path: str) -> str:
    """将章节列表写入 PDF 文件（使用 PyMuPDF）

    使用 CID-Font (China-Hans) 确保中文字符正确嵌入和提取。
    """
    # 尝试加载中文字体
    try:
        font = fitz.Font("china-s")
    except Exception:
        try:
            font = fitz.Font("china-hans")
        except Exception:
            font = None

    pdf = fitz.open()
    for sec in sections:
        page = pdf.new_page()
        rect = page.rect
        # 标题
        page.insert_text(
            (72, 72), sec["title"], fontsize=14,
            fontname="china-s" if font else "helv",
        )
        # 正文
        y = 100
        for para in sec["content"].split("\n"):
            if y > rect.height - 60:
                page = pdf.new_page()
                y = 72
            page.insert_text(
            (72, y), para.strip(), fontsize=11, fontname="china-s" if font else "helv")
            y += 20
    pdf.save(path)
    pdf.close()
    return path


# ═══════════════════════════════════════════════════════════════
# 公共 API
# ═══════════════════════════════════════════════════════════════

def generate_valid_bidding_doc(
    filepath: Optional[str] = None,
    fmt: str = "pdf",
) -> str:
    """
    生成一份完整的、合规的招标文件用于测试。

    包含：
    - 8 个章节（含全部 5 大必备章节）
    - 所有关键术语完整（公开招标、评审标准、投标截止、投票有效期等）
    - 无禁用词、无排他性条款

    Args:
        filepath: 输出路径（默认自动生成临时文件）
        fmt: 格式 "pdf" 或 "docx"

    Returns:
        生成的文件路径
    """
    if filepath is None:
        filepath = tempfile.NamedTemporaryFile(suffix=f".{fmt}", delete=False).name

    if fmt == "docx":
        _write_docx(COMPLIANT_SECTIONS, filepath)
    else:
        _write_pdf(COMPLIANT_SECTIONS, filepath)

    return filepath


def generate_invalid_bidding_doc(
    violations: list[str],
    filepath: Optional[str] = None,
    fmt: str = "pdf",
) -> str:
    """
    生成一份包含指定违规的招标文件用于测试。

    Args:
        violations: 违规类型列表，可选值：
            - "missing_section": 缺失章节（只保留 2 个章节）
            - "forbidden_words":  包含禁用词（指定品牌、本地注册等）
            - "missing_keywords": 缺少关键字（移除评审标准、投标截止等）
            - "all":              同时包含上述所有违规
        filepath: 输出路径（默认自动生成临时文件）
        fmt: 格式 "pdf" 或 "docx"

    Returns:
        生成的文件路径
    """
    if filepath is None:
        filepath = tempfile.NamedTemporaryFile(suffix=f".{fmt}", delete=False).name

    sections = list(COMPLIANT_SECTIONS)

    if "all" in violations or "missing_section" in violations:
        # 只保留招标公告和投标须知，其余章节全部缺失
        sections = [s for s in sections if s["title"] in (
            "第一章 招标公告",
            "第五章 投标人须知",
        )]

    if "all" in violations or "forbidden_words" in violations:
        # 在资格要求章节插入禁用词
        for s in sections:
            if "资格要求" in s["title"]:
                s["content"] += "\n" + FORBIDDEN_WORDS_CONTENT

    if "all" in violations or "missing_keywords" in violations:
        # 替换各章节内容为缺少关键字的版本
        for s in sections:
            if "评审办法" in s["title"]:
                s["content"] = MISSING_KEYWORDS_CONTENT
            if "投标人须知" in s["title"]:
                s["content"] = MISSING_KEYWORDS_CONTENT
            if "招标公告" in s["title"]:
                s["content"] = MISSING_KEYWORDS_CONTENT

    if fmt == "docx":
        _write_docx(sections, filepath)
    else:
        _write_pdf(sections, filepath)

    return filepath


def generate_bidding_doc_for_industry(
    industry: str,
    filepath: Optional[str] = None,
    fmt: str = "pdf",
) -> str:
    """
    生成适应特定行业的招标文件。

    Args:
        industry: "construction" / "it" / "healthcare"
        filepath: 输出路径
        fmt: "pdf" 或 "docx"

    Returns:
        生成的文件路径
    """
    import copy
    sections = copy.deepcopy(COMPLIANT_SECTIONS)

    industry_content = {
        "it": {
            "第二章 招标范围": (
                "技术参数要求：CPU主频≥2.0GHz，内存≥16GB。\n"
                "系统集成要求：须对接现有OA系统。\n"
                "信息安全等级保护：系统须达到等保2.0三级要求。\n"
                "国产化要求：须适配国产CPU和操作系统。\n"
                "数据安全：须满足数据加密和隐私保护要求。"
            ),
        },
        "construction": {
            "第三章 投标人资格要求": (
                "1. 投标人须具有建筑工程施工总承包二级及以上资质。\n"
                "2. 投标人须具有有效的安全生产许可证。\n"
                "3. 项目经理须持有建筑工程专业一级注册建造师证书。\n"
                "4. 投标人须提供近五年类似项目业绩至少3个。\n"
                "5. 本工程接受联合体投标。\n"
                "6. 禁止围标串标。"
            ),
            "第四章 评审办法": (
                "本项目采用综合评分法。\n"
                "评审标准：技术方案35分，价格30分，业绩20分，服务15分。\n"
                "工程质量标准：达到国家现行施工验收规范合格标准。"
            ),
        },
        "healthcare": {
            "第三章 投标人资格要求": (
                "1. 投标人须提供有效的医疗器械注册证。\n"
                "2. 投标人须具有医疗器械经营许可证。\n"
                "3. 投标产品须经国家药品监督管理局批准注册。\n"
                "4. 投标人须提供完善的售后服务和培训方案。\n"
                "5. 投标人须负责设备的安装调试和计量校准。"
            ),
        },
    }

    if filepath is None:
        filepath = tempfile.NamedTemporaryFile(suffix=f".{fmt}", delete=False).name

    if industry in industry_content:
        for title, content in industry_content[industry].items():
            for s in sections:
                if s["title"] == title:
                    s["content"] = content

    if fmt == "docx":
        _write_docx(sections, filepath)
    else:
        _write_pdf(sections, filepath)

    return filepath


def generate_doc_with_page_count(
    page_count: int,
    filepath: Optional[str] = None,
) -> str:
    """
    生成指定页数的 PDF 文件用于分页测试。

    Args:
        page_count: 目标页数
        filepath: 输出路径

    Returns:
        生成的文件路径
    """
    if filepath is None:
        filepath = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name

    try:
        font = fitz.Font("china-s")
    except Exception:
        try:
            font = fitz.Font("china-hans")
        except Exception:
            font = None

    pdf = fitz.open()
    for i in range(page_count):
        page = pdf.new_page()
        page.insert_text(
            (72, 720), f"第 {i+1} 页 / 共 {page_count} 页", fontsize=11,
            fontname="china-s" if font else "helv",
        )
        page.insert_text(
            (72, 700), "招标文件测试内容。" * 30, fontsize=9,
            fontname="china-s" if font else "helv",
        )
    pdf.save(filepath)
    pdf.close()
    return filepath
