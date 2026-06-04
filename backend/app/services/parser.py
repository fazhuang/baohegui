"""文档解析服务 - PDF/Word 章节抽取与结构化

支持 PDF（PyMuPDF，结构化+字体感知）和 Word（python-docx Heading 样式优先）格式的招标文件解析。

改进要点 v3:
1. Word 解析: Heading 样式优先识别章节结构，样式不存在时回退文本正则检测。
2. PDF 解析: 位置+字体感知，通过字体统计检测标题（比正文大/粗的文本行）。
3. 层级支持: 多级标题识别（一级/二级/三级），在 SectionInfo.level 中体现。
4. 模糊章节类型匹配: bi-gram 相似度匹配，处理标题变体。
5. 表格提取: Word 文档表格内容以 [TABLE: ...] 形式嵌入正文。
6. 编号列表过滤: 区分章节级标记（第X章）和列表级标记（1. 2. （1）），避免正文编号误检。
7. TOC 页面检测: 目录页跳过标题检测，避免目录条目被误认为章节。
8. 水印过滤: 长数字 ID（>10位）被过滤，粗体检测使用正确的标志位 (flags & 16)。
"""

from __future__ import annotations

import logging
import os
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import docx  # python-docx
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# ── 五大必备章节关键字映射 ──────────────────────────────────────
# 每个标准化类型用一个判别式关键词 + 补充同义词列表。
# 判别式关键词用于模糊相似度匹配的种子，同义词用于精确匹配。
SECTION_SEED_KEYWORDS: dict[str, str] = {
    "招标公告": "招标公告",
    "招标范围": "招标范围",
    "资格要求": "资格要求",
    "评审办法": "评审办法",
    "投标须知": "投标须知",
    "合同条款": "合同条款",
    "投标文件格式": "投标文件格式",
    "技术要求": "技术要求",
    "投标保证金": "投标保证金",
    "报价要求": "报价要求",
    "履约要求": "履约要求",
    "保密条款": "保密条款",
    "知识产权": "知识产权",
}

REQUIRED_SECTION_MAP: dict[str, list[str]] = {
    "招标公告": ["招标公告", "投标邀请", "投标邀请书", "竞争性谈判公告", "采购公告", "招标公告"],
    "招标范围": ["项目概况与招标范围", "项目概述与招标范围", "招标范围", "采购范围",
              "建设规模", "工程概况"],
    "资格要求": ["投标人资格要求", "申请人资格要求", "资格要求", "投标人资格",
              "供应商资格", "合格投标人", "投标人资质", "资质要求"],
    "评审办法": ["评审办法", "评标办法", "综合评分法", "评分标准", "评分细则",
              "评审标准", "评标标准", "评审方法", "评标方法"],
    "投标须知": ["投标人须知前附表", "投标人须知及前附表", "投标须知", "投标人须知",
              "投标说明", "投标人须知与合同", "投标人须知"],
    "合同条款": ["合同条款", "合同主要条款", "合同文本", "合同草案", "采购合同",
              "合同条款及格式", "合同格式"],
    "投标文件格式": ["投标文件格式", "投标文件组成", "投标文件编制", "投标文件的编制",
                 "电子投标文件的格式", "电子投标文件格式"],
    "投标保证金": ["投标保证金"],
    "技术要求": ["技术参数", "技术要求", "技术规格", "技术标准", "技术需求",
              "采购需求", "货物需求一览表", "技术标准和要求", "工程量清单",
              "采购项目需求", "项目需求"],
    "报价要求": ["报价要求", "报价说明", "报价方式", "报价一览表", "投标报价"],
    "履约要求": ["履约要求", "项目实施方案", "实施计划", "履约能力", "项目组织"],
    "保密条款": ["保密条款", "保密要求", "保密承诺", "信息安全", "商业秘密"],
    "知识产权": ["知识产权", "知识产权归属", "专利归属", "技术成果归属", "知识产权条款"],
}

# ── 标题格式正则（拆分为高/低置信度两组）───────────────────────

# CHAPTER_PATTERNS: 章节级标记，高置信度，不需要额外验证即可视为标题
# NOTE: "一、" patterns are moved to LIST_PATTERNS because they appear
# frequently in contract clauses (一、保修范围, 二、违约责任, etc.)
CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十\d]+[章节节部分]\s*\S"),
    re.compile(r"^第[一二三四五六七八九十\d]+条\s*\S"),
]

# LIST_PATTERNS: 列表级标记，低置信度，仅在包含已知章节关键词时才视为标题
# NOTE: "第X条" patterns are moved to CHAPTER_PATTERNS to avoid false
# positives from legal-clause numbering in engineering bid documents
LIST_PATTERNS = [
    re.compile(r"^\d+[、\.\s]\s*\S"),
    re.compile(r"^[A-Z]+[\.\s]\s*\S"),
    re.compile(r"^[一二三四五六七八九十]+[、\.\s]\s*\S"),
    re.compile(r"^[（(][一二三四五六七八九十\d]+[)）]\s*\S"),  # （一）/ (一)
]

# 完整集合（向后兼容/其他用途）
HEADING_PATTERNS = CHAPTER_PATTERNS + LIST_PATTERNS

# 已知章节同义词列表（用于 _is_heading_line 第一关检测）
# 只包含 >= 4 个汉字的短语，避免正文中的短词误检
ALL_SECTION_PATTERNS = [kw for pats in REQUIRED_SECTION_MAP.values() for kw in pats if len(kw) >= 4]


# ── PDF 字号统计阈值 ──────────────────────────────────────
PDF_BODY_FONT_MARGIN = 0.10       # 正文浮动 10% 仍算正文
PDF_HEADING_SCORE_THRESHOLD = 0.5 # 标题判定阈值
PDF_MIN_HEADING_SIZE = 11.0       # 略小于最小字号，过滤装饰

# ── 模糊匹配阈值 ───────────────────────────────────────────
SECTION_MATCH_MIN_RATIO = 0.45    # 最小重叠比
SECTION_MATCH_MIN_CHARS = 2       # 至少匹配 2 个字符


@dataclass
class TextSpan:
    """PDF 文本片段（带字体元信息）"""
    text: str
    font_size: float
    is_bold: bool
    font_name: str


@dataclass
class PageLine:
    """带字体/位置元信息的文本行"""
    text: str
    page_num: int
    bbox: tuple[float, float, float, float] = (0, 0, 0, 0)
    spans: list[TextSpan] = field(default_factory=list)
    max_font_size: float = 0.0
    is_bold: bool = False
    heading_score: float = 0.0


@dataclass
class SectionInfo:
    """文档章节信息"""

    title: str
    section_type: str  # 标准化章节类型
    content: str = ""
    level: int = 1  # 章节层级（1=一级标题，2=二级标题……）
    page_start: int = 0
    page_end: int = 0
    headings: list[str] = field(default_factory=list)  # 子标题列表

    def add_content(self, text: str) -> None:
        self.content += text

    def append_subheading(self, heading_title: str) -> None:
        self.headings.append(heading_title)

    def content_length(self) -> int:
        return len(self.content)

    def to_dict(self) -> dict:
        return {
            "section_type": self.section_type,
            "title": self.title,
            "content_length": self.content_length(),
            "content_preview": self.content[:500],
            "page_start": self.page_start,
            "page_end": self.page_end,
            "level": self.level,
            "headings": self.headings[:20],
        }


@dataclass
class ParsedDocument:
    """解析后的文档结构"""

    filename: str
    page_count: int = 0
    full_text: str = ""
    sections: dict[str, str] = field(default_factory=dict)  # section_type -> content
    raw_sections: list[SectionInfo] = field(default_factory=list)
    headings: list[str] = field(default_factory=list)  # 全文标题列表

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "page_count": self.page_count,
            "full_text_length": len(self.full_text),
            "sections": {
                sec_type: content[:200] for sec_type, content in self.sections.items()
            },
            "raw_sections": [s.to_dict() for s in self.raw_sections],
            "headings": self.headings[:50],
        }

    def get_section_count(self) -> int:
        """获取识别到的章节数量"""
        return len(self.sections)

    def get_missing_required_sections(self) -> list[str]:
        """返回缺失的必备章节列表"""
        found = set(self.sections.keys())
        required = {"招标公告", "招标范围", "资格要求", "评审办法", "投标须知"}
        return list(required - found)


class DocumentParserError(Exception):
    """文档解析异常"""

    pass


# ═══════════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════════

def _char_bigram_set(text: str) -> set[str]:
    """计算字符级 bi-gram 集合，用于相似度比较"""
    return {text[i:i+2] for i in range(len(text) - 1)}


def _bigram_overlap(a: str, b: str) -> float:
    """计算两个字符串的 bi-gram 重叠比"""
    if not a or not b:
        return 0.0
    bigrams_a = _char_bigram_set(a)
    bigrams_b = _char_bigram_set(b)
    if not bigrams_a or not bigrams_b:
        return 0.0
    intersection = bigrams_a & bigrams_b
    return len(intersection) / max(len(bigrams_a), len(bigrams_b))


def _is_toc_page(page_text: str) -> bool:
    """Detect if a page is a Table of Contents (目录) page.

    Criteria:
    - Contains the word '目录' AND
    - Has 3+ chapter-level patterns (第X章), indicating it's a listing page, not content
    """
    has_mulu = '目录' in page_text
    if not has_mulu:
        return False
    chapter_count = len(re.findall(r'第[一二三四五六七八九十\d]+章', page_text))
    return chapter_count >= 3


def _fuzzy_match_section_type(title: str) -> Optional[str]:
    """
    基于 bi-gram 重叠比进行模糊章节类型匹配。

    三步策略:
    1. 精确子串匹配（原方式，保留反向匹配）
    2. 每个类型判别式关键词的 bi-gram 重叠比 >= 阈值
    3. 每个类型同义词列表的 bi-gram 重叠比 >= 阈值
    优先级: 1 > 2 > 3，取匹配度最高的类型。
    """
    # Step 1: exact substring match
    for sec_type, patterns in REQUIRED_SECTION_MAP.items():
        for pattern in patterns:
            # Pattern is inside title (e.g., "资格要求" in "第一章 资格要求")
            if pattern in title:
                return sec_type
            # Title is inside pattern but only when title is substantial enough
            if title in pattern and len(title) >= len(pattern) * 0.6:
                return sec_type

    # Normalize title for fuzzy matching: remove heading prefixes
    title_clean = re.sub(
        r'^第[一二三四五六七八九十\d]+[章节节部分]|^[一二三四五六七八九十]+[、\.\s]|^\d+[、\.\s]|^[A-Z]+[\.\s]|^[（(][^)）]+[)）]',
        '',
        title,
    ).strip()
    # Remove parenthetical suffixes and trailing noise
    title_clean = re.sub(r'[（(].*[)）]|[:：].*$', '', title_clean).strip()

    if not title_clean or len(title_clean) < SECTION_MATCH_MIN_CHARS:
        return None

    # Step 2 & 3: bi-gram overlap scoring
    best_type: Optional[str] = None
    best_score = 0.0

    for sec_type, patterns in REQUIRED_SECTION_MAP.items():
        seed = SECTION_SEED_KEYWORDS.get(sec_type, sec_type)
        seed_score = _bigram_overlap(title_clean, seed)
        patterns_score = max(
            (_bigram_overlap(title_clean, p) for p in patterns if len(p) >= SECTION_MATCH_MIN_CHARS),
            default=0.0,
        )
        score = max(seed_score, patterns_score)
        if score > best_score:
            best_score = score
            best_type = sec_type

    if best_score >= SECTION_MATCH_MIN_RATIO:
        return best_type
    return None


def _compute_pdf_font_stats(doc: fitz.Document) -> dict[int, float]:
    """
    收集 PDF 每页字号分布，返回 (page_num -> body_font_size) 映射。

    每页正文体字号 = 该页出现频次最高的字号；
    如果某页无数据，使用全文众数。
    封面页（第0页）从全局统计中排除，避免大标题拉偏全局 body_size。
    """
    page_stats: dict[int, Counter] = {}
    all_sizes: Counter = Counter()

    for page_num in range(len(doc)):
        page = doc[page_num]
        blocks = page.get_text('dict')['blocks']
        sizes: list[float] = []
        for block in blocks:
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                for span in line.get('spans', []):
                    size = span.get('size', 0)
                    if size >= PDF_MIN_HEADING_SIZE:
                        sizes.append(size)
        page_counter: Counter = Counter()
        for s in sizes:
            key = round(s * 2) / 2
            page_counter[key] += 1
            # Exclude page 0 (cover) from global stats
            if page_num > 0:
                all_sizes[key] += 1
        page_stats[page_num] = page_counter

    global_mode = all_sizes.most_common(1)
    global_body_size = float(global_mode[0][0]) if global_mode else 12.0
    result: dict[int, float] = {}
    for page_num in range(len(doc)):
        counter = page_stats.get(page_num, Counter())
        if counter:
            page_mode = counter.most_common(1)[0][0]
            result[page_num] = float(page_mode)
        else:
            result[page_num] = global_body_size
    return result


def _compute_heading_scores(doc: fitz.Document, body_sizes: dict[int, float]) -> list[PageLine]:
    """
    对 PDF 每行文本计算标题得分，返回 PageLine 列表。

    标题得分基于:
    - 字号相对正文大小的比值（越大越像标题）
    - 是否加粗（加粗加分）
    - 文本长度（过长的行降低标题可能性）
    """
    page_lines: list[PageLine] = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        body_size = body_sizes.get(page_num, 12.0)
        blocks = page.get_text('dict')['blocks']
        for block in blocks:
            if block.get('type') != 0:
                continue
            for line_data in block.get('lines', []):
                line_bbox = line_data.get('bbox', (0, 0, 0, 0))
                spans_data = line_data.get('spans', [])
                if not spans_data:
                    continue
                full_text = ''
                font_sizes: list[float] = []
                is_bold = False
                span_list: list[TextSpan] = []
                for span in spans_data:
                    text = span.get('text', '').strip()
                    if not text:
                        continue
                    # Filter watermark: long pure-numeric IDs (>10 digits), common in GPC PDFs
                    if len(text) > 10 and text.isdigit():
                        continue
                    size = span.get('size', 0)
                    flags = span.get('flags', 0)
                    font = span.get('font', '')
                    bold = bool(flags & 16)  # bit 4 = bold (not bit 1 = italic)
                    full_text += text
                    font_sizes.append(size)
                    is_bold = is_bold or bold
                    span_list.append(TextSpan(text=text, font_size=size, is_bold=bold, font_name=font))
                if not full_text.strip():
                    continue
                max_size = max(font_sizes) if font_sizes else 0
                score = 0.0
                if max_size > 0 and body_size > 0:
                    size_ratio = max_size / body_size
                    if size_ratio > 1.0 + PDF_BODY_FONT_MARGIN:
                        score = min((size_ratio - 1.0) * 3.0, 0.8)
                if is_bold:
                    score += 0.15
                text_len = len(full_text.strip())
                if text_len <= 5:
                    score += 0.15
                elif text_len >= 60:
                    score -= 0.2
                page_lines.append(PageLine(
                    text=full_text.strip(),
                    page_num=page_num + 1,
                    bbox=line_bbox,
                    spans=span_list,
                    max_font_size=max_size,
                    is_bold=is_bold,
                    heading_score=score,
                ))
    return page_lines


def _is_probable_heading(page_line: PageLine) -> bool:
    """基于字体+文本长度判断是否为标题

    增强策略：
    - 字号评分 >= 阈值 → 高置信度标题
    - 中文字体感知：在字号相近时，只有已知关键词的行才视为标题
    - 短行 + 含中文 + 已知关键词 → 独立章节标题（无编号模式，如「评审办法」）
    - 短行 + 加粗 + 含中文 → 可能为标题（无编号加粗模式）
    """
    if page_line.heading_score >= PDF_HEADING_SCORE_THRESHOLD:
        return True
    stripped = page_line.text.strip()
    if not stripped or len(stripped) > 60:
        return False
    # Only check chapter-level patterns (not list patterns) when font score is low
    for pat in CHAPTER_PATTERNS:
        if pat.match(stripped):
            return True
    # For list patterns, require known section keywords when font score is low
    has_keyword = any(p in stripped for p in ALL_SECTION_PATTERNS if len(p) >= 4)
    if has_keyword:
        for pat in LIST_PATTERNS:
            if pat.match(stripped):
                return True
    # ── 增强：无编号独立关键词标题 ──────────────────────────
    # 短行（≤20字）+ 包含已知章节关键词 + 中文 → 独立章节标题
    # 处理「评审办法」「技术要求」等无编号、无列表标记的独立标题行
    if has_keyword and len(stripped) <= 20:
        has_chinese = bool(re.search(r'[一-鿿]', stripped))
        if has_chinese:
            return True
    # ── 增强：加粗短行含中文 → 潜在标题 ─────────────────────
    # 当字体评分略低但明确加粗、短行、含中文时，标记为标题
    # 处理仅靠加粗（而非大字号）标记的章节标题
    if page_line.is_bold and len(stripped) <= 30 and page_line.heading_score >= 0.15:
        has_chinese = bool(re.search(r'[一-鿿]', stripped))
        if has_chinese:
            return True
    return has_keyword and not stripped.rstrip("。！？；：,.!").isdigit()


def _estimate_heading_level(page_line: PageLine, body_size: float) -> int:
    """根据字号比估算标题层级（1=最顶层）"""
    if body_size <= 0 or page_line.max_font_size <= 0:
        return 1
    ratio = page_line.max_font_size / body_size
    if ratio >= 1.5:
        return 1
    elif ratio >= 1.25:
        return 2
    else:
        return 3


def _docx_heading_level(style_name: str) -> int:
    """从 Word 样式名提取标题层级"""
    match = re.search(r'(?i)heading\s*(\d+)', style_name)
    if match:
        return int(match.group(1))
    return 1


class DocumentParser:
    """文档解析器：支持 PDF 和 Word，自动识别章节并抽取结构化内容"""

    # ── 章节检测 ────────────────────────────────────────────

    @staticmethod
    def _detect_section_type(title: str) -> Optional[str]:
        """
        根据标题文本检测标准化章节类型。

        优先精确匹配（子串/反向），回退到模糊 bi-gram 匹配。
        """
        return _fuzzy_match_section_type(title)

    @staticmethod
    def _is_heading_line(line: str) -> Optional[str]:
        """
        根据文本内容判断是否为章节标题。

        Detection strategies (strict to lenient):
        1. CHAPTER_PATTERNS (第X章, 一、, 第X条) -> high confidence chapter headings
        2. LIST_PATTERNS (1., 2., (1)) -> only if containing known section keywords
        3. Pure keyword match (short lines <=12 chars, must contain Chinese)
        4. Fuzzy bi-gram match for unknown but heading-like text
        """
        stripped = line.strip()
        if not stripped or len(stripped) > 70:
            return None

        has_known_keyword = any(p in stripped for p in ALL_SECTION_PATTERNS)

        # Strategy A: CHAPTER_PATTERNS (chapter-level markers, high confidence)
        for pat in CHAPTER_PATTERNS:
            if not pat.match(stripped):
                continue
            caption = re.sub(
                r"^(第[^章节部分]+[章节节部分]|[一二三四五六七八九十\d]+[、\.\s]|第[一二三四五六七八九十\d]+条)\s*",
                "",
                stripped,
            )
            if not caption or len(caption) > 30:
                continue
            if has_known_keyword and any(caption.startswith(p) for p in ALL_SECTION_PATTERNS):
                return stripped
            if len(caption) <= 8 and has_known_keyword:
                return stripped
            break

        # Strategy B: LIST_PATTERNS (list-level markers, needs keyword verification)
        # Numbered items like "1. xxx", "A. xxx", "(1) xxx" are only
        # treated as headings when they contain known section keywords
        for pat in LIST_PATTERNS:
            if pat.match(stripped):
                if has_known_keyword:
                    caption = re.sub(
                        r"^(\d+[、\.\s]|[A-Z]+[\.\s]|[（(][^)）]+[)）])\s*",
                        "",
                        stripped,
                    )
                    if caption and any(p in caption for p in ALL_SECTION_PATTERNS):
                        return stripped
                # List items without known keywords are NOT headings
                return None

        # Strategy C: pure keyword match (short lines with known keywords, must contain Chinese)
        # Increased from 12 to 20 chars to catch headings like "技术标准和要求" (7 chars)
        # Require Chinese characters to filter out pure numeric references like "2.1", "5.1"
        if has_known_keyword and len(stripped) <= 20:
            stripped_clean = stripped.rstrip("。！？；：,.!")
            has_chinese = bool(re.search(r'[一-鿿]', stripped_clean))
            if has_chinese:
                for p in ALL_SECTION_PATTERNS:
                    if stripped == p or stripped_clean == p:
                        return stripped
                # Also match if the line is short and contains a known keyword
                # as a standalone heading (e.g. "第三章 评审办法")
                return stripped

        # Strategy D: fuzzy heading match via bi-gram (for unknown heading-like text)
        if len(stripped) <= 40 and not has_known_keyword:
            section_type = _fuzzy_match_section_type(stripped)
            if section_type is not None:
                return stripped

        return None

    # ── 核心抽取逻辑 ────────────────────────────────────────

    @staticmethod
    def _extract_sections_from_lines(
        lines: list[tuple[int, str]],
        toc_pages: set[int] | None = None,
    ) -> tuple[dict[str, str], list[SectionInfo], list[str]]:
        """
        从带页码的文本行列表中提取章节结构。

        Args:
            lines: 每元素为 (page_number, text_line)
            toc_pages: 已知 TOC 页面集合，这些页面的行不触发标题检测

        Returns:
            (sections_dict, raw_sections_list, headings_list)
        """
        if toc_pages is None:
            toc_pages = set()

        sections: dict[str, str] = {}
        all_headings: list[str] = []
        raw_sections: list[SectionInfo] = []

        current: Optional[SectionInfo] = None

        for page_num, line in lines:
            stripped = line.strip()
            if not stripped:
                if current:
                    current.add_content("\n")
                continue

            # Skip heading detection on TOC pages
            heading_text = None
            if page_num not in toc_pages:
                heading_text = DocumentParser._is_heading_line(stripped)

            if heading_text:
                all_headings.append(heading_text)
                sec_type = DocumentParser._detect_section_type(heading_text) or heading_text

                if current:
                    if sec_type == current.section_type:
                        current.add_content(f"\n【{heading_text}】\n")
                        continue
                    prev = current.content
                    if current.section_type in sections:
                        sections[current.section_type] += "\n" + prev
                    else:
                        sections[current.section_type] = prev
                    if current not in raw_sections:
                        raw_sections.append(current)

                current = SectionInfo(
                    title=heading_text,
                    section_type=sec_type,
                    content="",
                    level=1,
                    page_start=page_num,
                    page_end=page_num,
                )
            elif current:
                current.add_content(stripped + "\n")
                current.page_end = page_num

        if current and current.content.strip():
            sections[current.section_type] = current.content
            if current not in raw_sections:
                raw_sections.append(current)

        return sections, raw_sections, all_headings

    # ── PDF 结构化解析（字体感知）───────────────────────────

    @staticmethod
    def _extract_sections_from_page_lines(
        page_lines: list[PageLine],
        body_sizes: dict[int, float] | None = None,
        toc_pages: set[int] | None = None,
    ) -> tuple[dict[str, str], list[SectionInfo], list[str]]:
        """
        从 PageLine 列表中提取章节结构（PDF 增强版）。

        利用 heading_score 和字体信息进行更精确的标题检测。

        Args:
            page_lines: 带字体/位置信息的文本行列表
            body_sizes: 每页的正文字号
            toc_pages: 已知 TOC 页面集合，跳过标题检测

        Returns:
            (sections_dict, raw_sections_list, headings_list)
        """
        if toc_pages is None:
            toc_pages = set()

        sections: dict[str, str] = {}
        all_headings: list[str] = []
        raw_sections: list[SectionInfo] = []

        current: Optional[SectionInfo] = None

        for pl in page_lines:
            stripped = pl.text.strip()
            if not stripped:
                continue

            # Skip heading detection on TOC pages
            is_heading = False
            if pl.page_num not in toc_pages:
                is_heading = _is_probable_heading(pl)

            if is_heading:
                all_headings.append(stripped)
                sec_type = DocumentParser._detect_section_type(stripped) or stripped

                body_size = 12.0
                if body_sizes is not None:
                    body_size = body_sizes.get(pl.page_num - 1, 12.0)
                level = _estimate_heading_level(pl, body_size)

                if current and sec_type == current.section_type and level >= current.level:
                    current.add_content(f"\n【{stripped}】\n")
                    current.append_subheading(stripped)
                    continue

                if current and current.content.strip():
                    prev_content = current.content.rstrip()
                    if current.section_type in sections:
                        sections[current.section_type] += "\n\n" + prev_content
                    else:
                        sections[current.section_type] = prev_content
                    if current not in raw_sections:
                        raw_sections.append(current)

                current = SectionInfo(
                    title=stripped,
                    section_type=sec_type,
                    content="",
                    level=level,
                    page_start=pl.page_num,
                    page_end=pl.page_num,
                )
            elif current:
                current.add_content(stripped + "\n")
                current.page_end = pl.page_num

        if current and current.content.strip():
            final_content = current.content.rstrip()
            if current.section_type in sections:
                sections[current.section_type] += "\n\n" + final_content
            else:
                sections[current.section_type] = final_content
            if current not in raw_sections:
                raw_sections.append(current)

        return sections, raw_sections, all_headings

    # ── PDF 解析 ────────────────────────────────────────────

    def _parse_pdf_structured(self, doc: fitz.Document, filepath: str) -> ParsedDocument:
        """
        基于字体感知的 PDF 结构化解析。

        流程:
        1. 检测 TOC 页面
        2. 统计每页字号分布 -> 确定正文大小
        3. 每行计算标题得分
        4. 基于标题得分检测章节结构（跳过 TOC 页）
        5. 回退到纯文本检测如果字体分析效果不佳
        """
        # Detect TOC pages (pages with '目录' + multiple chapter references)
        toc_pages: set[int] = set()
        for page_num in range(min(5, len(doc))):  # only check first 5 pages
            page_text = doc[page_num].get_text()
            if _is_toc_page(page_text):
                toc_pages.add(page_num + 1)  # 1-indexed
                logger.info("Detected TOC page %d", page_num + 1)

        body_sizes = _compute_pdf_font_stats(doc)
        page_lines = _compute_heading_scores(doc, body_sizes)

        detected_headings = sum(1 for pl in page_lines if pl.heading_score >= PDF_HEADING_SCORE_THRESHOLD)

        # ── 双向解析融合：字体分析 + 文本检测取最优 ──────────────
        # 即使字体分析检测到足够标题，也运行文本检测作为补充
        # 两种方法的结果合并，取章节覆盖更全的一方
        font_sections, font_raw_sections, font_headings = {}, [], []
        text_sections, text_raw_sections, text_headings = {}, [], []

        if detected_headings >= 3:
            font_sections, font_raw_sections, font_headings = self._extract_sections_from_page_lines(
                page_lines, body_sizes, toc_pages,
            )
            logger.info(
                "PDF 字体分析检测到 %d 个标题候选，识别 %d 个章节",
                detected_headings, len(font_sections),
            )

        # 始终运行文本检测作为兜底/补充
        simple_lines = [(pl.page_num, pl.text) for pl in page_lines]
        text_sections, text_raw_sections, text_headings = self._extract_sections_from_lines(
            simple_lines, toc_pages,
        )
        logger.info(
            "PDF 文本检测识别到 %d 个章节",
            len(text_sections),
        )

        # ── 择优合并：取章节覆盖更全的结果 ──────────────────────
        # 字体分析更精确但可能遗漏无特殊格式的标题；
        # 文本检测覆盖面更广但可能有误检。
        # 策略：以字体分析结果为基础，补充文本检测中识别到但字体分析遗漏的章节
        if len(font_sections) >= len(text_sections):
            sections = font_sections
            raw_sections = font_raw_sections
            headings = font_headings
            # 补充文本检测中发现的额外章节（字体分析遗漏的）
            for sec_type, sec_content in text_sections.items():
                if sec_type not in sections:
                    sections[sec_type] = sec_content
                    logger.info(
                        "文本检测补充字体分析遗漏的章节: [%s] (%d 字符)",
                        sec_type, len(sec_content),
                    )
        else:
            sections = text_sections
            raw_sections = text_raw_sections
            headings = text_headings
            # 补充字体分析中发现的额外章节
            for sec_type, sec_content in font_sections.items():
                if sec_type not in sections:
                    sections[sec_type] = sec_content
                    logger.info(
                        "字体分析补充文本检测遗漏的章节: [%s] (%d 字符)",
                        sec_type, len(sec_content),
                    )

        if not sections:
            logger.warning(
                "PDF 解析未识别到任何章节！字体候选=%d, 文本候选=%d。"
                "文件可能为扫描件（需OCR）或格式特殊。",
                len(font_sections), len(text_sections),
            )

        page_count = len(doc)
        full_text = "\n".join(pl.text for pl in page_lines)

        return ParsedDocument(
            filename=Path(filepath).name,
            page_count=page_count,
            full_text=full_text,
            sections=sections,
            raw_sections=raw_sections,
            headings=headings,
        )

    def parse_pdf(self, filepath: str) -> ParsedDocument:
        """
        解析 PDF 文件（字体感知 + 位置分析）。

        使用 PyMuPDF get_text("dict") 获取字体大小/粗体信息，
        通过字号统计自动识别标题和正文。

        Args:
            filepath: PDF 文件路径

        Returns:
            ParsedDocument
        """
        if not os.path.exists(filepath):
            raise DocumentParserError(f"文件不存在: {filepath}")

        try:
            doc = fitz.open(filepath)
        except Exception as e:
            raise DocumentParserError(f"PyMuPDF 无法打开文件: {filepath}: {e}") from e

        try:
            result = self._parse_pdf_structured(doc, filepath)
        finally:
            doc.close()

        return result

    # ── Word 解析 ───────────────────────────────────────────

    # ── Word 结构化解析 ────────────────────────────────────

    @staticmethod
    def _extract_docx_tables(doc: docx.Document) -> list[str]:
        """提取 Word 文档中的表格内容为可读文本"""
        table_texts: list[str] = []
        for table_idx, table in enumerate(doc.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append(f"[TABLE {table_idx + 1}]\n" + "\n".join(rows))
        return table_texts

    def _parse_docx_by_style(self, doc: docx.Document, filepath: str) -> ParsedDocument:
        """
        利用 Word Heading 样式识别章节。

        创建 Heading 样式的层级树，将非 Heading 段落归入最近标题下。
        表格内容提取为 [TABLE: ...] 文本块。
        """
        sections: dict[str, str] = {}
        all_headings: list[str] = []
        raw_sections: list[SectionInfo] = []
        current: Optional[SectionInfo] = None
        total_paras = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            total_paras += 1
            style_name = (para.style.name or "").lower() if para.style else ""
            is_style_heading = "heading" in style_name

            if is_style_heading and text:
                level = _docx_heading_level(para.style.name)
                sec_type = self._detect_section_type(text) or text
                all_headings.append(text)

                if current and current.content.strip():
                    final_content = current.content.rstrip()
                    if current.section_type in sections:
                        sections[current.section_type] += "\n\n" + final_content
                    else:
                        sections[current.section_type] = final_content
                    if current not in raw_sections:
                        raw_sections.append(current)

                current = SectionInfo(
                    title=text,
                    section_type=sec_type,
                    content="",
                    level=level,
                    page_start=total_paras,
                    page_end=total_paras,
                )
            elif current and text:
                current.add_content(text + "\n")
                current.page_end = total_paras

        if current and current.content.strip():
            final_content = current.content.rstrip()
            if current.section_type in sections:
                sections[current.section_type] += "\n\n" + final_content
            else:
                sections[current.section_type] = final_content
            if current not in raw_sections:
                raw_sections.append(current)

        # 提取表格并归属到最近章节（而非全部追加到每个章节）
        table_texts = self._extract_docx_tables(doc)

        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if table_texts:
            full_text += "\n\n" + "\n\n".join(table_texts)
            # 将表格追加到最后一个章节（而非全部章节）
            if raw_sections:
                raw_sections[-1].add_content("\n\n" + "\n\n".join(table_texts))
            elif current and current.content.strip():
                current.add_content("\n\n" + "\n\n".join(table_texts))

        sections_updated: dict[str, str] = {}
        for rs in raw_sections:
            sections_updated[rs.section_type] = sections.get(rs.section_type, rs.content)

        return ParsedDocument(
            filename=Path(filepath).name,
            page_count=max(total_paras // 30 + 1, 1),
            full_text=full_text,
            sections=sections or sections_updated,
            raw_sections=raw_sections,
            headings=all_headings,
        )

    def _parse_docx_text_based(self, doc: docx.Document, filepath: str) -> ParsedDocument:
        """原始文本检测方式（保留原行为）"""
        lines: list[tuple[int, str]] = []
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            lines.append((idx // 30 + 1, text))

        full_text = "\n".join(l for _, l in lines)
        sections, raw_sections, headings = self._extract_sections_from_lines(lines)
        page_count = max(len(doc.paragraphs) // 30 + 1, 1)

        # Also extract tables — append to last section only
        table_texts = DocumentParser._extract_docx_tables(doc)
        if table_texts:
            full_text += "\n\n" + "\n\n".join(table_texts)
            if raw_sections:
                raw_sections[-1].add_content("\n\n" + "\n\n".join(table_texts))

        return ParsedDocument(
            filename=Path(filepath).name,
            page_count=page_count,
            full_text=full_text,
            sections=sections,
            raw_sections=raw_sections,
            headings=headings,
        )

    def _parse_docx_structured(self, doc: docx.Document, filepath: str) -> ParsedDocument:
        """
        基于 Word Heading 样式进行结构化解析。

        优先使用 Heading 1/2/3 样式识别章节结构；
        样式层级映射到 SectionInfo.level；
        表格内容以 [TABLE: ...] 形式嵌入。
        """
        heading_count = sum(
            1 for p in doc.paragraphs
            if p.style and p.style.name and "heading" in p.style.name.lower()
        )

        if heading_count >= 3:
            return self._parse_docx_by_style(doc, filepath)
        else:
            result = self._parse_docx_text_based(doc, filepath)
            if len(result.sections) >= 1:
                return result
            return result

    def parse_docx(self, filepath: str) -> ParsedDocument:
        """
        解析 Word (.docx) 文件（Heading 样式优先）。

        优先利用 python-docx 的 Heading 1/2/3 样式识别章节结构，
        样式不存在时回退到文本正则检测。
        自动提取表格内容。

        Args:
            filepath: .docx 文件路径

        Returns:
            ParsedDocument
        """
        if not os.path.exists(filepath):
            raise DocumentParserError(f"文件不存在: {filepath}")

        try:
            doc = docx.Document(filepath)
        except Exception as e:
            raise DocumentParserError(f"python-docx 无法打开文件: {filepath}: {e}") from e

        return self._parse_docx_structured(doc, filepath)

    # ── 统一入口 ────────────────────────────────────────────

    def parse(self, filepath: str) -> ParsedDocument:
        """
        根据文件扩展名自动选择解析方式。

        Args:
            filepath: 文件路径

        Returns:
            ParsedDocument

        Raises:
            DocumentParserError: 不支持的文件格式或解析失败
        """
        ext = Path(filepath).suffix.lower()
        if ext == ".pdf":
            return self.parse_pdf(filepath)
        elif ext == ".docx":
            return self.parse_docx(filepath)
        else:
            raise DocumentParserError(f"不支持的文件格式 '{ext}'，仅支持 PDF 和 .docx")


# 模块级便捷实例
parser = DocumentParser()
